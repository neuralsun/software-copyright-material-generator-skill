#!/usr/bin/env python3
"""Build a Chinese software-copyright application-form DOCX from JSON.

The renderer intentionally separates product facts from legal identity data.  It
never guesses an applicant, copyright owner, credential number, address, or
contact detail.  Missing legal fields remain blank in the DOCX and make the
machine-readable readiness report return ``ready: false``.

The bundled form is a layout template, not an official online submission form.
Always transfer or verify the final values in the current registration portal.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import tempfile
import zipfile
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = SCRIPT_DIR.parent / "assets" / "application-form-template.docx"
SCHEMA_VERSION = 1
EXPECTED_ROWS = 37
EXPECTED_COLUMNS = 13
RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
PLACEHOLDER_RE = re.compile(
    r"(?:TODO|TBD|FIXME|待填|待补|待定|后期添加|后期补充|示例名称|XXX|某某)", re.I
)


class ApplicationFormError(ValueError):
    """Raised when a configuration or reference template is invalid."""


def _norm(value: Any) -> str:
    return re.sub(r"\s+", "", "" if value is None else str(value))


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (list, tuple)):
        return "\n".join(str(item).strip() for item in value if str(item).strip())
    return str(value).strip()


def _mapping(value: Any, label: str) -> dict[str, Any]:
    if value is None:
        return {}
    if not isinstance(value, dict):
        raise ApplicationFormError(f"{label} 必须是 JSON 对象")
    return value


def _get_path(data: dict[str, Any], dotted: str) -> Any:
    current: Any = data
    for part in dotted.split("."):
        if not isinstance(current, dict) or part not in current:
            return None
        current = current[part]
    return current


def _pick(data: dict[str, Any], *paths: str, default: Any = None) -> Any:
    for path in paths:
        value = _get_path(data, path)
        if value is not None:
            return value
    return default


def _required_text(data: dict[str, Any], label: str, *paths: str) -> str:
    value = _text(_pick(data, *paths))
    if not value:
        raise ApplicationFormError(f"缺少必填配置：{label}（{paths[0]}）")
    if PLACEHOLDER_RE.search(value):
        raise ApplicationFormError(f"{label} 含有占位文字，不能写入申请表：{value!r}")
    return value


def _parse_date(value: Any, label: str, *, optional: bool = False) -> str:
    raw = _text(value)
    if not raw and optional:
        return ""
    if not raw:
        raise ApplicationFormError(f"缺少必填配置：{label}")
    compact = raw.replace("年", "-").replace("月", "-").replace("日", "")
    compact = compact.replace("/", "-").replace(".", "-")
    try:
        parsed = date.fromisoformat(compact)
    except ValueError as exc:
        raise ApplicationFormError(f"{label} 必须是有效日期（YYYY-MM-DD）：{raw!r}") from exc
    return f"{parsed.year:04d}年{parsed.month:02d}月{parsed.day:02d}日"


def _normalize_version(value: str) -> str:
    cleaned = value.strip()
    if re.fullmatch(r"\d+(?:\.\d+){1,3}", cleaned):
        cleaned = "V" + cleaned
    if not re.fullmatch(r"[Vv]\d+(?:\.\d+){1,3}(?:[-+._A-Za-z0-9]*)?", cleaned):
        raise ApplicationFormError(
            "software.version 应采用用户确认的登记版本格式；新申请默认使用 V1.0"
        )
    return "V" + cleaned[1:]


def _visible_char_count(value: str) -> int:
    return len(re.sub(r"\s+", "", value))


def _coerce_positive_int(value: Any, label: str) -> int:
    if isinstance(value, bool):
        raise ApplicationFormError(f"{label} 必须是正整数")
    try:
        result = int(value)
    except (TypeError, ValueError) as exc:
        raise ApplicationFormError(f"{label} 必须是正整数") from exc
    if result <= 0:
        raise ApplicationFormError(f"{label} 必须大于 0")
    return result


def _as_bool(value: Any, label: str) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str) and value.strip().lower() in {"true", "yes", "1", "是"}:
        return True
    if isinstance(value, str) and value.strip().lower() in {"false", "no", "0", "否"}:
        return False
    raise ApplicationFormError(f"{label} 必须是 JSON 布尔值 true/false")


def _load_config(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
    except FileNotFoundError as exc:
        raise ApplicationFormError(f"配置文件不存在：{path}") from exc
    except json.JSONDecodeError as exc:
        raise ApplicationFormError(
            f"配置不是有效 JSON：{path}（第 {exc.lineno} 行第 {exc.colno} 列）"
        ) from exc
    if not isinstance(data, dict):
        raise ApplicationFormError("配置根节点必须是 JSON 对象")
    return data


def validate_template(document: Document, source: Path | str) -> dict[str, int]:
    """Strictly validate the invariant layout used by the application renderer."""
    if len(document.tables) != 1:
        raise ApplicationFormError(
            f"申请表模板必须且只能包含 1 个正文表格；{source} 实际为 {len(document.tables)} 个"
        )
    table = document.tables[0]
    rows = len(table.rows)
    columns = len(table.columns)
    grid = table._tbl.tblGrid
    grid_columns = len(grid.gridCol_lst) if grid is not None else 0
    if rows != EXPECTED_ROWS or columns != EXPECTED_COLUMNS or grid_columns != EXPECTED_COLUMNS:
        raise ApplicationFormError(
            "申请表模板结构必须为 37 行、13 个逻辑列；"
            f"{source} 实际为 {rows} 行、{columns} 列、{grid_columns} 个网格列"
        )

    anchors = (
        (0, 0, "计算机软件著作权登记申请表"),
        (1, 0, "软件名称"),
        (7, 0, "开发方式"),
        (15, 0, "权利范围"),
        (16, 0, "软件功能和技术特点"),
        (24, 1, "登记软件的主要功能和技术特点"),
        (26, 0, "申请人信息"),
        (34, 0, "软件鉴别材料"),
    )
    for row, column, expected in anchors:
        actual = _norm(table.cell(row, column).text)
        if _norm(expected) not in actual:
            raise ApplicationFormError(
                f"模板锚点校验失败：第 {row + 1} 行第 {column + 1} 列应包含 {expected!r}"
            )
    return {"tables": 1, "rows": rows, "columns": columns}


def _clear_paragraph(paragraph: Any) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _format_run(run: Any, *, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "宋体")


def _set_cell(
    cell: Any,
    value: str,
    *,
    required: bool = False,
    size: float = 9.0,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    justify: bool = False,
) -> None:
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if value:
        run = paragraph.add_run(value)
        _format_run(run, size=size, color=RED if required else BLACK, bold=bold)


def _set_field_run(paragraph: Any, instruction: str, *, size: float = 9.0) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    _format_run(run, size=size, color=BLACK)
    run._r.extend((begin, instr, separate, placeholder, end))


def _replace_page_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        for table in list(footer.tables):
            footer._element.remove(table._tbl)
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        _clear_paragraph(paragraph)
        for extra in list(footer.paragraphs[1:]):
            footer._element.remove(extra._p)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("第 ")
        _format_run(run, size=9, color=BLACK)
        _set_field_run(paragraph, "PAGE")
        run = paragraph.add_run(" 页 / 共 ")
        _format_run(run, size=9, color=BLACK)
        _set_field_run(paragraph, "NUMPAGES")
        run = paragraph.add_run(" 页")
        _format_run(run, size=9, color=BLACK)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _clean_core_properties(document: Document, title: str) -> None:
    props = document.core_properties
    now = datetime.now(timezone.utc).replace(tzinfo=None, microsecond=0)
    props.title = title
    props.subject = "计算机软件著作权登记申请表"
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.keywords = ""
    props.category = ""
    props.content_status = ""
    props.identifier = ""
    props.language = "zh-CN"
    props.version = ""
    props.revision = 1
    props.created = now
    props.modified = now
    # python-docx does not accept ``None`` for this optional datetime; remove
    # the inherited element instead of preserving the source document's print
    # timestamp.
    last_printed = props._element.find(qn("cp:lastPrinted"))
    if last_printed is not None:
        props._element.remove(last_printed)


def _sanitize_package_metadata(path: Path) -> None:
    """Remove custom properties and blank identity-bearing extended metadata."""
    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent))
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            temp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as target:
            for item in source.infolist():
                if item.filename == "docProps/custom.xml":
                    continue
                payload = source.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    root = ET.fromstring(payload)
                    for child in list(root):
                        if child.attrib.get("PartName") == "/docProps/custom.xml":
                            root.remove(child)
                    payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                elif item.filename == "_rels/.rels":
                    root = ET.fromstring(payload)
                    for child in list(root):
                        if child.attrib.get("Type", "").endswith("/custom-properties"):
                            root.remove(child)
                    payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                elif item.filename == "docProps/app.xml":
                    root = ET.fromstring(payload)
                    for child in root.iter():
                        local = child.tag.rsplit("}", 1)[-1]
                        if local in {"Company", "Manager", "HyperlinkBase", "Template"}:
                            child.text = ""
                        elif local == "TotalTime":
                            child.text = "0"
                    payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                target.writestr(item, payload)
        os.replace(temp_path, path)
    finally:
        temp_path.unlink(missing_ok=True)


def _save_document(document: Document, output: Path, title: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(output.parent))
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        _clean_core_properties(document, title)
        _replace_page_footer(document)
        document.save(temp_path)
        _sanitize_package_metadata(temp_path)
        os.replace(temp_path, output)
    finally:
        temp_path.unlink(missing_ok=True)


def _genericize_template(document: Document) -> None:
    table = document.tables[0]
    for row, column in ((1, 4), (1, 12), (2, 4), (2, 12), (5, 4)):
        _set_cell(table.cell(row, column), "")
    _set_cell(table.cell(3, 2), "□ 原创", size=9)
    _set_cell(
        table.cell(4, 2),
        "□ 修改（含翻译软件、合成软件和升级版本软件）\n"
        "  修改软件须经原权利人授权\n"
        "  □ 原有软件已经登记\n"
        "    原登记号：\n"
        "    修改（翻译或合成）软件作品说明：",
        size=8.5,
    )
    _set_cell(table.cell(5, 12), "    年    月    日\n（未发表则不填）", size=8.5)
    _set_cell(table.cell(6, 12), "国家：        城市：", size=8.5)
    _set_cell(
        table.cell(7, 4),
        "□ 独立开发      □ 合作开发      □ 委托开发      □ 下达任务开发",
        size=9,
    )
    for row in (9, 10):
        for column in (3, 8, 11):
            _set_cell(table.cell(row, column), "")
    for row in (13, 14):
        for column in (3, 7, 11):
            _set_cell(table.cell(row, column), "")
    _set_cell(table.cell(11, 3), "□ 继承          □ 受让          □ 承受", size=9)
    _set_cell(table.cell(15, 3), "□ 全部      □ 部分权利（                              ）", size=9)
    for row in (18, 20, 22, 23, 25):
        _set_cell(table.cell(row, 1), "")
    for row in range(26, 34):
        _set_cell(table.cell(row, 5), "")
    for row, label in (
        (26, "电话："),
        (27, "邮政编码："),
        (28, "手机："),
        (29, "传真："),
        (30, "电话："),
        (31, "邮政编码："),
        (32, "手机："),
        (33, "传真："),
    ):
        _set_cell(table.cell(row, 10), label, size=9)
    _set_cell(table.cell(34, 1), "□ 一般交存", size=9)
    _set_cell(table.cell(35, 1), "□ 例外交存", size=9)


def _normalize_choice(value: Any, label: str, choices: dict[str, str]) -> str:
    raw = _text(value).lower().replace(" ", "")
    if raw in choices:
        return choices[raw]
    allowed = "、".join(sorted(set(choices.values())))
    raise ApplicationFormError(f"{label} 无效；可选值：{allowed}")


def _legal_value(mapping: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = _text(mapping.get(key))
        if value:
            return "" if PLACEHOLDER_RE.search(value) else value
    return ""


def _collect_legal_data(config: dict[str, Any]) -> tuple[list[dict[str, str]], dict[str, str], dict[str, str], list[str]]:
    raw_owners = _pick(config, "copyright_owners", "rights_holders", default=[])
    if isinstance(raw_owners, dict):
        raw_owners = [raw_owners]
    if raw_owners is None:
        raw_owners = []
    if not isinstance(raw_owners, list):
        raise ApplicationFormError("copyright_owners 必须是 JSON 数组")
    if len(raw_owners) > 2:
        raise ApplicationFormError("本模板最多容纳 2 名原始著作权人")

    owners: list[dict[str, str]] = []
    missing: list[str] = []
    if not raw_owners:
        missing.append("copyright_owners[0]")
    for index, raw in enumerate(raw_owners):
        source = _mapping(raw, f"copyright_owners[{index}]")
        owner = {
            "name": _legal_value(source, "name"),
            "credential_no": _legal_value(source, "credential_no", "certificate_no", "id_no"),
            "nationality": _legal_value(source, "nationality"),
            "address": _legal_value(source, "address"),
        }
        for field in ("name", "credential_no", "nationality", "address"):
            if not owner[field]:
                missing.append(f"copyright_owners[{index}].{field}")
        owners.append(owner)

    applicant_source = _mapping(config.get("applicant"), "applicant")
    applicant = {
        "name": _legal_value(applicant_source, "name"),
        "address": _legal_value(applicant_source, "address", "actual_address"),
        "postal_code": _legal_value(applicant_source, "postal_code", "postcode"),
        "contact": _legal_value(applicant_source, "contact", "contact_person"),
        "phone": _legal_value(applicant_source, "phone", "telephone"),
        "mobile": _legal_value(applicant_source, "mobile", "mobile_phone"),
        "email": _legal_value(applicant_source, "email", "e_mail"),
        "fax": _legal_value(applicant_source, "fax"),
    }
    for field in ("name", "address", "postal_code", "contact", "mobile", "email"):
        if not applicant[field]:
            missing.append(f"applicant.{field}")

    agent_source = _mapping(config.get("agent"), "agent")
    agent = {
        "name": _legal_value(agent_source, "name"),
        "address": _legal_value(agent_source, "address"),
        "postal_code": _legal_value(agent_source, "postal_code", "postcode"),
        "contact": _legal_value(agent_source, "contact", "contact_person"),
        "phone": _legal_value(agent_source, "phone", "telephone"),
        "mobile": _legal_value(agent_source, "mobile", "mobile_phone"),
        "email": _legal_value(agent_source, "email", "e_mail"),
        "fax": _legal_value(agent_source, "fax"),
    }
    if any(agent.values()):
        for field in ("name", "address", "postal_code", "contact", "mobile", "email"):
            if not agent[field]:
                missing.append(f"agent.{field}")
    return owners, applicant, agent, missing


def _fill_contact_rows(table: Any, start_row: int, party: dict[str, str]) -> None:
    values = (party["name"], party["address"], party["contact"], party["email"])
    trailing = (
        f"电话：{party['phone']}",
        f"邮政编码：{party['postal_code']}",
        f"手机：{party['mobile']}",
        f"传真：{party['fax']}",
    )
    for offset, value in enumerate(values):
        _set_cell(table.cell(start_row + offset, 5), value, required=bool(value), size=8.5)
    for offset, value in enumerate(trailing):
        _set_cell(table.cell(start_row + offset, 10), value, required=bool(value.split("：", 1)[1]), size=8.5)


def _render_form(document: Document, config: dict[str, Any]) -> dict[str, Any]:
    table = document.tables[0]
    full_name = _required_text(
        config, "软件全称", "software.full_name", "software.name", "software_name"
    )
    short_name = _text(_pick(config, "software.short_name", "software_short_name"))
    classification_no = _text(_pick(config, "software.classification_no", "classification_no"))
    version = _normalize_version(
        _required_text(config, "软件版本", "software.version", "software_version", "version")
    )
    completion_date = _parse_date(
        _pick(config, "software.completion_date", "completion_date"), "software.completion_date"
    )
    original = _as_bool(_pick(config, "software.original", "original"), "software.original")
    if not original:
        raise ApplicationFormError("本生成器仅用于原创软件；software.original 必须为 true")

    development_mode = _normalize_choice(
        _pick(config, "software.development_mode", "development_mode"),
        "software.development_mode",
        {
            "independent": "独立开发", "独立开发": "独立开发",
            "collaborative": "合作开发", "cooperative": "合作开发", "合作开发": "合作开发",
            "commissioned": "委托开发", "委托开发": "委托开发",
            "assigned": "下达任务开发", "task": "下达任务开发", "下达任务开发": "下达任务开发",
        },
    )
    raw_rights = _pick(config, "software.rights_scope", "rights_scope")
    rights_details = ""
    if isinstance(raw_rights, dict):
        rights_mode = _normalize_choice(
            raw_rights.get("mode"), "software.rights_scope.mode",
            {"all": "全部", "全部": "全部", "partial": "部分权利", "部分权利": "部分权利"},
        )
        rights_details = _text(raw_rights.get("details"))
    else:
        rights_mode = _normalize_choice(
            raw_rights, "software.rights_scope",
            {"all": "全部", "全部": "全部", "partial": "部分权利", "部分权利": "部分权利"},
        )
    if rights_mode == "部分权利" and not rights_details:
        raise ApplicationFormError("选择部分权利时必须填写 software.rights_scope.details")

    deposit_mode = _normalize_choice(
        _pick(config, "deposit.mode", "deposit_mode", default="general"),
        "deposit.mode",
        {"general": "一般交存", "一般交存": "一般交存"},
    )
    hardware = _required_text(
        config, "硬件环境", "software.hardware_environment", "hardware_environment"
    )
    software_environment = _required_text(
        config, "软件环境", "software.software_environment", "software_environment"
    )
    languages = _required_text(
        config, "编程语言及版本", "software.programming_languages", "programming_languages"
    )
    source_lines = _coerce_positive_int(
        _pick(
            config,
            "software.source_line_count",
            "software.program_nonblank_lines",
            "source_line_count",
            "program_nonblank_lines",
        ),
        "software.source_line_count",
    )
    feature_text = _required_text(
        config,
        "主要功能和技术特点",
        "software.function_and_technical_features",
        "function_and_technical_features",
    )
    feature_chars = _visible_char_count(feature_text)
    if not 500 <= feature_chars <= 1000:
        raise ApplicationFormError(
            "software.function_and_technical_features 去除空白后的长度必须为 500–1000 字；"
            f"当前为 {feature_chars} 字"
        )

    first_publication = _parse_date(
        _pick(config, "software.first_publication_date", "first_publication.date"),
        "software.first_publication_date",
        optional=True,
    )
    publication_country = _text(
        _pick(config, "software.first_publication_country", "first_publication.country")
    )
    publication_city = _text(
        _pick(config, "software.first_publication_city", "first_publication.city")
    )
    if first_publication and (not publication_country or not publication_city):
        raise ApplicationFormError("首次发表时必须同时填写发表国家和城市")
    if not first_publication and (publication_country or publication_city):
        raise ApplicationFormError("未填写首次发表日期时，不得单独填写发表国家或城市")

    owners, applicant, agent, missing_legal = _collect_legal_data(config)

    _set_cell(table.cell(1, 4), full_name, required=True, size=10, bold=True)
    _set_cell(table.cell(1, 12), classification_no, required=bool(classification_no), size=9)
    _set_cell(table.cell(2, 4), short_name, required=bool(short_name), size=9)
    _set_cell(table.cell(2, 12), version, required=True, size=10, bold=True)
    _set_cell(table.cell(3, 2), "☑ 原创", required=True, size=9)
    _set_cell(
        table.cell(4, 2),
        "□ 修改（含翻译软件、合成软件和升级版本软件）\n"
        "  修改软件须经原权利人授权\n"
        "  □ 原有软件已经登记\n"
        "    原登记号：\n"
        "    修改（翻译或合成）软件作品说明：",
        size=8.5,
    )
    _set_cell(table.cell(5, 4), completion_date, required=True, size=9)
    if first_publication:
        _set_cell(table.cell(5, 12), first_publication, required=True, size=9)
        _set_cell(
            table.cell(6, 12),
            f"国家：{publication_country}    城市：{publication_city}",
            required=True,
            size=9,
        )
    else:
        _set_cell(table.cell(5, 12), "    年    月    日\n（未发表则不填）", size=8.5)
        _set_cell(table.cell(6, 12), "国家：        城市：", size=8.5)

    modes = ("独立开发", "合作开发", "委托开发", "下达任务开发")
    development_text = "      ".join(
        ("☑ " if value == development_mode else "□ ") + value for value in modes
    )
    _set_cell(table.cell(7, 4), development_text, required=True, size=9)

    for row in (9, 10):
        for column in (3, 8, 11):
            _set_cell(table.cell(row, column), "")
    for index, owner in enumerate(owners):
        row = 9 + index
        owner_identity = owner["name"]
        if owner["credential_no"]:
            owner_identity += f"\n证件号码：{owner['credential_no']}"
        _set_cell(table.cell(row, 3), owner_identity, required=bool(owner_identity), size=8)
        _set_cell(table.cell(row, 8), owner["nationality"], required=bool(owner["nationality"]), size=8.5)
        _set_cell(table.cell(row, 11), owner["address"], required=bool(owner["address"]), size=8)

    _set_cell(table.cell(11, 3), "□ 继承          □ 受让          □ 承受", size=9)
    for row in (13, 14):
        for column in (3, 7, 11):
            _set_cell(table.cell(row, column), "")
    if rights_mode == "全部":
        rights_text = "☑ 全部      □ 部分权利（                              ）"
    else:
        rights_text = f"□ 全部      ☑ 部分权利（{rights_details}）"
    _set_cell(table.cell(15, 3), rights_text, required=True, size=9)

    _set_cell(table.cell(18, 1), hardware, required=True, size=8.5)
    _set_cell(table.cell(20, 1), software_environment, required=True, size=8.5)
    _set_cell(table.cell(22, 1), languages, required=True, size=8.5)
    _set_cell(table.cell(23, 1), f"程序量：代码 {source_lines} 条", required=True, size=9)
    _set_cell(table.cell(25, 1), feature_text, required=True, size=8.5, justify=True)

    _fill_contact_rows(table, 26, applicant)
    _fill_contact_rows(table, 30, agent)
    _set_cell(table.cell(34, 1), "☑ 一般交存", required=True, size=9)
    _set_cell(table.cell(35, 1), "□ 例外交存", size=9)

    return {
        "software_name": full_name,
        "version": version,
        "completion_date": completion_date,
        "development_mode": development_mode,
        "rights_scope": rights_mode,
        "deposit_mode": deposit_mode,
        "program_nonblank_lines": source_lines,
        "feature_visible_characters": feature_chars,
        "published": bool(first_publication),
        "missing_legal_fields": missing_legal,
    }


def create_generic_template(source: Path, destination: Path) -> dict[str, Any]:
    if source.resolve() == destination.resolve():
        raise ApplicationFormError("模板源文件与输出文件不能是同一路径")
    try:
        document = Document(source)
    except Exception as exc:
        raise ApplicationFormError(f"无法打开模板源文件：{source}（{exc}）") from exc
    layout = validate_template(document, source)
    _genericize_template(document)
    _save_document(document, destination, "计算机软件著作权登记申请表模板")
    check = Document(destination)
    validate_template(check, destination)
    # Ensure the template itself has no product data in the principal value cells.
    table = check.tables[0]
    for row, column in ((1, 4), (2, 12), (5, 4), (18, 1), (20, 1), (22, 1), (23, 1), (25, 1)):
        if _norm(table.cell(row, column).text):
            raise ApplicationFormError(
                f"通用模板清理失败：第 {row + 1} 行第 {column + 1} 列仍有产品数据"
            )
    return {"template": str(destination), "layout": layout}


def build_application_form(
    config_path: Path,
    template_path: Path,
    output_path: Path,
    report_path: Path,
) -> dict[str, Any]:
    if template_path.resolve() == output_path.resolve():
        raise ApplicationFormError("模板文件与申请表输出文件不能是同一路径")
    config = _load_config(config_path)
    try:
        document = Document(template_path)
    except Exception as exc:
        raise ApplicationFormError(f"无法打开申请表模板：{template_path}（{exc}）") from exc
    layout = validate_template(document, template_path)
    facts = _render_form(document, config)
    _save_document(document, output_path, f"{facts['software_name']}_软件著作权登记申请表")

    check = Document(output_path)
    output_layout = validate_template(check, output_path)
    body_text = "\n".join(cell.text for table in check.tables for row in table.rows for cell in row.cells)
    if facts["software_name"] not in body_text or facts["version"] not in body_text:
        raise ApplicationFormError("输出复核失败：软件名称或版本未写入申请表")

    missing_legal = facts.pop("missing_legal_fields")
    report = {
        "schema_version": SCHEMA_VERSION,
        "artifact": "software-copyright-application-form",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "ready": not missing_legal,
        "ready_definition": "申请表配置数据完整；仍须由申请人核验、按受理平台要求签章并在线提交。",
        "missing_legal_fields": missing_legal,
        "manual_actions": [
            "由申请人核验著作权人、证件号码、地址及联系人信息，严禁使用示例或虚构身份。",
            "按当前登记平台要求在线填写或导入，并由有权主体完成签章及签章日期。",
            "确认开发完成日期、首次发表状态、软件名称和版本与说明书、代码材料完全一致。",
        ],
        "input": {"config": str(config_path.resolve()), "template": str(template_path.resolve())},
        "output": str(output_path.resolve()),
        "report": str(report_path.resolve()),
        "template_validation": layout,
        "output_validation": output_layout,
        "form_facts": facts,
    }
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="由通用 JSON 配置生成计算机软件著作权登记申请表 DOCX，并输出就绪性报告。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "生成申请表：\n"
            "  python build_application_form.py --config application.json --output 申请表.docx\n\n"
            "从已核验的 37 行模板制作去产品化资产（维护用途）：\n"
            "  python build_application_form.py --create-template-from 参考申请表.docx "
            "--template-output ../assets/application-form-template.docx\n\n"
            "缺少法律身份字段时仍生成留白申请表，但报告 ready=false；"
            "加 --require-ready 可令该情况返回非零退出码。"
        ),
    )
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--config", type=Path, help="UTF-8 JSON 配置文件")
    mode.add_argument("--create-template-from", type=Path, metavar="DOCX", help="从参考表生成通用模板")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help=f"申请表模板（默认：{DEFAULT_TEMPLATE}）")
    parser.add_argument("--output", type=Path, help="申请表 DOCX 输出路径")
    parser.add_argument("--report", type=Path, help="就绪性 JSON 报告；默认与申请表同名并加 .report.json")
    parser.add_argument("--require-ready", action="store_true", help="法律字段不完整（ready=false）时返回退出码 3")
    parser.add_argument("--template-output", type=Path, default=DEFAULT_TEMPLATE, help="--create-template-from 的输出路径")
    return parser


def main(argv: Iterable[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(list(argv) if argv is not None else None)
    try:
        if args.create_template_from:
            result = create_generic_template(args.create_template_from.resolve(), args.template_output.resolve())
            print(json.dumps({"ok": True, **result}, ensure_ascii=False))
            return 0
        if args.output is None:
            parser.error("使用 --config 时必须同时提供 --output")
        output = args.output.resolve()
        if output.suffix.lower() != ".docx":
            raise ApplicationFormError("--output 必须以 .docx 结尾")
        report = (args.report or output.with_suffix(".report.json")).resolve()
        result = build_application_form(args.config.resolve(), args.template.resolve(), output, report)
        print(
            json.dumps(
                {
                    "ok": True,
                    "ready": result["ready"],
                    "output": result["output"],
                    "report": result["report"],
                    "missing_legal_fields": len(result["missing_legal_fields"]),
                },
                ensure_ascii=False,
            )
        )
        return 3 if args.require_ready and not result["ready"] else 0
    except ApplicationFormError as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        return 2


if __name__ == "__main__":
    sys.exit(main())
