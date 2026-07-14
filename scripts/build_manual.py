#!/usr/bin/env python3
"""Build a software design and user manual DOCX from a JSON specification.

The renderer is deliberately product-agnostic.  Software facts, prose, figures,
and tables must all come from the input JSON; this module only owns document
structure and typography.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DEFAULT_EAST_ASIA_FONT = "宋体"
DEFAULT_HEADING_FONT = "黑体"
DEFAULT_LATIN_FONT = "Times New Roman"
SUPPORTED_IMAGE_SUFFIXES = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff"}


class ManualConfigError(ValueError):
    """Raised when the JSON specification cannot produce a valid manual."""


@dataclass
class RenderState:
    figures: int = 0
    tables: int = 0
    paragraphs: int = 0
    headings: int = 0
    missing_images: int = 0


def _as_mapping(value: Any, label: str) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise ManualConfigError(f"{label} must be a JSON object")
    return value


def _as_list(value: Any, label: str) -> list[Any]:
    if value is None:
        return []
    if not isinstance(value, list):
        raise ManualConfigError(f"{label} must be a JSON array")
    return value


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    return str(value)


def _float(value: Any, default: float, label: str) -> float:
    if value is None:
        return default
    try:
        return float(value)
    except (TypeError, ValueError) as exc:
        raise ManualConfigError(f"{label} must be a number") from exc


def _int(value: Any, default: int, label: str) -> int:
    if value is None:
        return default
    try:
        return int(value)
    except (TypeError, ValueError) as exc:
        raise ManualConfigError(f"{label} must be an integer") from exc


def _rgb(value: Any, default: str = "000000") -> RGBColor:
    raw = _text(value).strip().lstrip("#") or default
    if not re.fullmatch(r"[0-9A-Fa-f]{6}", raw):
        raise ManualConfigError(f"invalid RGB color: {value!r}")
    return RGBColor.from_string(raw.upper())


def _get_or_add_fonts(r_pr: Any) -> Any:
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    return fonts


def set_run_font(
    run: Any,
    *,
    east_asia: str = DEFAULT_EAST_ASIA_FONT,
    latin: str = DEFAULT_LATIN_FONT,
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    fonts = _get_or_add_fonts(r_pr)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)


def set_style_font(
    style: Any,
    *,
    east_asia: str,
    latin: str,
    size: float,
    bold: bool = False,
) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    fonts = _get_or_add_fonts(r_pr)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)


def set_line_spacing(paragraph: Any, multiple: float) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(round(240 * multiple)))
    spacing.set(qn("w:lineRule"), "auto")


def set_cell_margins(cell: Any, *, top: int = 70, start: int = 100, bottom: int = 70, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table: Any, *, color: str = "A6A6A6", size: str = "4") -> None:
    table_pr = table._tbl.tblPr
    borders = table_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill.lstrip("#").upper())


def repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def clear_paragraph(paragraph: Any) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_field(paragraph: Any, instruction: str, display_text: str = "") -> Any:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_text, separate, text_node, end])
    return run


def enable_field_updates(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def restart_page_numbering(section: Any, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_section(section: Any, layout: dict[str, Any]) -> None:
    page = _as_mapping(layout.get("page", {}), "document.layout.page")
    margins = _as_mapping(layout.get("margins_cm", {}), "document.layout.margins_cm")
    section.page_width = Cm(_float(page.get("width_cm"), 21.0, "page.width_cm"))
    section.page_height = Cm(_float(page.get("height_cm"), 29.7, "page.height_cm"))
    section.top_margin = Cm(_float(margins.get("top"), 2.8, "margins_cm.top"))
    section.bottom_margin = Cm(_float(margins.get("bottom"), 2.5, "margins_cm.bottom"))
    section.left_margin = Cm(_float(margins.get("left"), 3.0, "margins_cm.left"))
    section.right_margin = Cm(_float(margins.get("right"), 2.6, "margins_cm.right"))
    section.header_distance = Cm(_float(layout.get("header_distance_cm"), 1.2, "header_distance_cm"))
    section.footer_distance = Cm(_float(layout.get("footer_distance_cm"), 1.2, "footer_distance_cm"))
    grid = section._sectPr.find(qn("w:docGrid"))
    if grid is None:
        grid = OxmlElement("w:docGrid")
        section._sectPr.append(grid)
    grid.set(qn("w:type"), "lines")
    grid.set(qn("w:linePitch"), str(_int(layout.get("line_pitch_twips"), 360, "line_pitch_twips")))


def prepare_styles(document: Document, style_cfg: dict[str, Any]) -> dict[str, Any]:
    body_font = _text(style_cfg.get("body_font")) or DEFAULT_EAST_ASIA_FONT
    heading_font = _text(style_cfg.get("heading_font")) or DEFAULT_HEADING_FONT
    latin_font = _text(style_cfg.get("latin_font")) or DEFAULT_LATIN_FONT
    body_size = _float(style_cfg.get("body_size_pt"), 10.5, "styles.body_size_pt")
    h1_size = _float(style_cfg.get("heading1_size_pt"), 18.0, "styles.heading1_size_pt")
    h2_size = _float(style_cfg.get("heading2_size_pt"), 15.0, "styles.heading2_size_pt")
    h3_size = _float(style_cfg.get("heading3_size_pt"), 12.0, "styles.heading3_size_pt")

    normal = document.styles["Normal"]
    set_style_font(normal, east_asia=body_font, latin=latin_font, size=body_size)
    normal.paragraph_format.space_after = Pt(_float(style_cfg.get("paragraph_after_pt"), 4.0, "styles.paragraph_after_pt"))

    for name, size in (("Heading 1", h1_size), ("Heading 2", h2_size), ("Heading 3", h3_size)):
        style = document.styles[name]
        set_style_font(style, east_asia=heading_font, latin=latin_font, size=size, bold=True)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0 if name == "Heading 1" else 10)

    try:
        caption = document.styles["Figure Caption"]
    except KeyError:
        caption = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(
        caption,
        east_asia=body_font,
        latin=latin_font,
        size=_float(style_cfg.get("caption_size_pt"), 12.0, "styles.caption_size_pt"),
        bold=bool(style_cfg.get("caption_bold", True)),
    )
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)

    return {
        "body_font": body_font,
        "heading_font": heading_font,
        "latin_font": latin_font,
        "body_size": body_size,
        "line_spacing": _float(style_cfg.get("line_spacing"), 1.25, "styles.line_spacing"),
        "first_line_indent_cm": _float(
            style_cfg.get("first_line_indent_cm"), 0.741, "styles.first_line_indent_cm"
        ),
        "caption_size": _float(style_cfg.get("caption_size_pt"), 12.0, "styles.caption_size_pt"),
        "caption_bold": bool(style_cfg.get("caption_bold", True)),
    }


def clean_core_properties(document: Document, software: dict[str, Any], document_cfg: dict[str, Any]) -> None:
    metadata = _as_mapping(document_cfg.get("metadata", {}), "document.metadata")
    name = _text(software.get("name")).strip()
    title_suffix = _text(document_cfg.get("title_suffix")).strip() or "软件设计与使用说明书"
    title_separator = _text(document_cfg.get("title_separator"))
    title = _text(document_cfg.get("title")).strip() or f"{name}{title_separator}{title_suffix}"
    now = datetime.now(timezone.utc)
    core = document.core_properties
    core.title = title
    core.subject = _text(metadata.get("subject")).strip() or "软件著作权登记说明书"
    core.author = _text(metadata.get("author")).strip()
    core.last_modified_by = _text(metadata.get("last_modified_by")).strip()
    core.comments = _text(metadata.get("comments")).strip()
    core.category = _text(metadata.get("category")).strip()
    core.keywords = _text(metadata.get("keywords")).strip()
    core.identifier = _text(metadata.get("identifier")).strip()
    core.language = _text(metadata.get("language")).strip() or "zh-CN"
    core.content_status = _text(metadata.get("content_status")).strip()
    core.created = now
    core.modified = now
    core.revision = 1


def add_body_paragraph(document: Document, block: dict[str, Any], style_info: dict[str, Any], state: RenderState) -> Any:
    text = _text(block.get("text"))
    paragraph = document.add_paragraph(style="Normal")
    alignment = _text(block.get("alignment")).strip().lower()
    alignments = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    if alignment:
        if alignment not in alignments:
            raise ManualConfigError(f"unknown paragraph alignment: {alignment}")
        paragraph.alignment = alignments[alignment]
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    indent = block.get("first_line_indent_cm", style_info["first_line_indent_cm"])
    if indent is not None:
        paragraph.paragraph_format.first_line_indent = Cm(_float(indent, 0.0, "paragraph.first_line_indent_cm"))
    paragraph.paragraph_format.space_before = Pt(_float(block.get("space_before_pt"), 0.0, "paragraph.space_before_pt"))
    paragraph.paragraph_format.space_after = Pt(_float(block.get("space_after_pt"), 4.0, "paragraph.space_after_pt"))
    paragraph.paragraph_format.keep_with_next = bool(block.get("keep_with_next", False))
    set_line_spacing(paragraph, _float(block.get("line_spacing"), style_info["line_spacing"], "paragraph.line_spacing"))
    run = paragraph.add_run(text)
    set_run_font(
        run,
        east_asia=_text(block.get("font")) or style_info["body_font"],
        latin=_text(block.get("latin_font")) or style_info["latin_font"],
        size=_float(block.get("size_pt"), style_info["body_size"], "paragraph.size_pt"),
        bold=bool(block.get("bold", False)),
        italic=bool(block.get("italic", False)),
        color=_rgb(block.get("color")) if block.get("color") else None,
    )
    state.paragraphs += 1
    return paragraph


def add_heading(document: Document, block: dict[str, Any], state: RenderState) -> Any:
    text = _text(block.get("text") or block.get("title")).strip()
    if not text:
        raise ManualConfigError("heading text cannot be empty")
    level = _int(block.get("level"), 1, "heading.level")
    if level not in (1, 2, 3):
        raise ManualConfigError("heading.level must be 1, 2, or 3")
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = bool(block.get("page_break_before", False))
    state.headings += 1
    return paragraph


def resolve_asset_path(raw_path: Any, base_dir: Path) -> Path:
    value = _text(raw_path).strip()
    if not value:
        raise ManualConfigError("figure.path cannot be empty")
    path = Path(value).expanduser()
    if not path.is_absolute():
        path = base_dir / path
    return path.resolve()


def add_figure(
    document: Document,
    block: dict[str, Any],
    *,
    base_dir: Path,
    max_width_cm: float,
    document_cfg: dict[str, Any],
    style_info: dict[str, Any],
    state: RenderState,
) -> None:
    path = resolve_asset_path(block.get("path"), base_dir)
    caption = _text(block.get("caption")).strip()
    if not caption and not bool(document_cfg.get("allow_missing_captions", False)):
        raise ManualConfigError(f"figure caption is required: {path}")
    if not caption:
        caption = path.stem
    if bool(document_cfg.get("auto_number_figures", True)) and not re.match(r"^图\s*\d+", caption):
        caption = f"图{state.figures + 1} {caption}"

    if not path.exists():
        if not bool(document_cfg.get("allow_missing_images", False)):
            raise FileNotFoundError(path)
        placeholder = document.add_paragraph()
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder.paragraph_format.keep_with_next = True
        run = placeholder.add_run(f"【后期补图：{caption}】")
        set_run_font(
            run,
            east_asia=style_info["body_font"],
            latin=style_info["latin_font"],
            size=style_info["body_size"],
        )
        state.missing_images += 1
    else:
        if path.suffix.lower() not in SUPPORTED_IMAGE_SUFFIXES:
            raise ManualConfigError(f"unsupported figure format for DOCX: {path.suffix}")
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(_float(block.get("space_before_pt"), 0.0, "figure.space_before_pt"))
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        requested_width = _float(block.get("width_cm"), min(15.35, max_width_cm), "figure.width_cm")
        width_cm = min(max(requested_width, 1.0), max_width_cm)
        shape = paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
        try:
            shape._inline.docPr.set("descr", caption)
            shape._inline.docPr.set("title", caption)
        except AttributeError:
            pass

    caption_paragraph = document.add_paragraph(style="Figure Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = None
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(_float(block.get("caption_after_pt"), 0.0, "figure.caption_after_pt"))
    set_line_spacing(caption_paragraph, _float(block.get("caption_line_spacing"), 1.4, "figure.caption_line_spacing"))
    run = caption_paragraph.add_run(caption)
    set_run_font(
        run,
        east_asia=style_info["body_font"],
        latin=style_info["latin_font"],
        size=style_info["caption_size"],
        bold=style_info["caption_bold"],
    )
    state.figures += 1


def add_table(
    document: Document,
    block: dict[str, Any],
    *,
    max_width_cm: float,
    style_info: dict[str, Any],
    state: RenderState,
) -> None:
    headers = [_text(item) for item in _as_list(block.get("headers"), "table.headers")]
    rows_raw = _as_list(block.get("rows"), "table.rows")
    rows: list[list[str]] = []
    for row_index, raw_row in enumerate(rows_raw, start=1):
        if not isinstance(raw_row, list):
            raise ManualConfigError(f"table row {row_index} must be an array")
        rows.append([_text(item) for item in raw_row])
    if not headers and not rows:
        raise ManualConfigError("table must contain headers or rows")
    column_count = len(headers) if headers else len(rows[0])
    if column_count < 1:
        raise ManualConfigError("table must contain at least one column")
    if headers and len(headers) != column_count:
        raise ManualConfigError("table header width mismatch")
    for row_index, row in enumerate(rows, start=1):
        if len(row) != column_count:
            raise ManualConfigError(f"table row {row_index} has {len(row)} cells; expected {column_count}")

    caption = _text(block.get("caption")).strip()
    if caption:
        if bool(block.get("auto_number", True)) and not re.match(r"^表\s*\d+", caption):
            caption = f"表{state.tables + 1} {caption}"
        caption_paragraph = document.add_paragraph(style="Figure Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.keep_with_next = True
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(
            caption_run,
            east_asia=style_info["body_font"],
            latin=style_info["latin_font"],
            size=style_info["caption_size"],
            bold=style_info["caption_bold"],
        )

    table = document.add_table(rows=1 if headers else 0, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=_text(block.get("border_color")) or "A6A6A6")
    widths_raw = _as_list(block.get("widths_cm"), "table.widths_cm")
    if widths_raw:
        if len(widths_raw) != column_count:
            raise ManualConfigError("table.widths_cm must match the number of columns")
        widths = [_float(value, 0.0, "table.widths_cm") for value in widths_raw]
        total = sum(widths)
        if total <= 0:
            raise ManualConfigError("table.widths_cm must have a positive total")
        if total > max_width_cm:
            scale = max_width_cm / total
            widths = [value * scale for value in widths]
    else:
        widths = [max_width_cm / column_count] * column_count

    if headers:
        repeat_table_header(table.rows[0])
        for index, value in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.width = Cm(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_shading(cell, _text(block.get("header_fill")) or "EAF2F8")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            set_line_spacing(paragraph, 1.2)
            run = paragraph.add_run(value)
            set_run_font(
                run,
                east_asia=style_info["body_font"],
                latin=style_info["latin_font"],
                size=style_info["body_size"],
                bold=True,
            )

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.width = Cm(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            set_line_spacing(paragraph, 1.2)
            run = paragraph.add_run(value)
            set_run_font(
                run,
                east_asia=style_info["body_font"],
                latin=style_info["latin_font"],
                size=style_info["body_size"],
            )
    state.tables += 1


def render_blocks(
    document: Document,
    blocks: Iterable[Any],
    *,
    base_dir: Path,
    max_width_cm: float,
    document_cfg: dict[str, Any],
    style_info: dict[str, Any],
    state: RenderState,
) -> None:
    for index, raw_block in enumerate(blocks, start=1):
        if isinstance(raw_block, str):
            block = {"type": "paragraph", "text": raw_block}
        elif isinstance(raw_block, dict):
            block = raw_block
        else:
            raise ManualConfigError(f"block {index} must be a string or object")
        block_type = _text(block.get("type")).strip().lower() or "paragraph"
        if block_type == "paragraph":
            add_body_paragraph(document, block, style_info, state)
        elif block_type == "heading":
            add_heading(document, block, state)
        elif block_type == "figure":
            add_figure(
                document,
                block,
                base_dir=base_dir,
                max_width_cm=max_width_cm,
                document_cfg=document_cfg,
                style_info=style_info,
                state=state,
            )
        elif block_type == "table":
            add_table(document, block, max_width_cm=max_width_cm, style_info=style_info, state=state)
        elif block_type == "page_break":
            document.add_page_break()
        else:
            raise ManualConfigError(f"unsupported block type: {block_type}")


def render_container(
    document: Document,
    container: dict[str, Any],
    *,
    heading_level: int | None,
    base_dir: Path,
    max_width_cm: float,
    document_cfg: dict[str, Any],
    style_info: dict[str, Any],
    state: RenderState,
) -> None:
    title = _text(container.get("title")).strip()
    if heading_level is not None:
        if not title:
            raise ManualConfigError("chapter or section title cannot be empty")
        add_heading(
            document,
            {
                "text": title,
                "level": _int(container.get("level"), heading_level, "container.level"),
                "page_break_before": bool(container.get("page_break_before", False)),
            },
            state,
        )
    paragraphs = _as_list(container.get("paragraphs"), "container.paragraphs")
    paragraph_blocks = [{"type": "paragraph", "text": _text(item)} for item in paragraphs]
    render_blocks(
        document,
        paragraph_blocks,
        base_dir=base_dir,
        max_width_cm=max_width_cm,
        document_cfg=document_cfg,
        style_info=style_info,
        state=state,
    )
    render_blocks(
        document,
        _as_list(container.get("blocks"), "container.blocks"),
        base_dir=base_dir,
        max_width_cm=max_width_cm,
        document_cfg=document_cfg,
        style_info=style_info,
        state=state,
    )
    child_level = 2 if heading_level is None else min(3, heading_level + 1)
    for section in _as_list(container.get("sections"), "container.sections"):
        render_container(
            document,
            _as_mapping(section, "section"),
            heading_level=child_level,
            base_dir=base_dir,
            max_width_cm=max_width_cm,
            document_cfg=document_cfg,
            style_info=style_info,
            state=state,
        )


def add_cover(document: Document, software: dict[str, Any], document_cfg: dict[str, Any], style_info: dict[str, Any]) -> None:
    name = _text(software.get("name")).strip()
    version = _text(software.get("version")).strip()
    title_suffix = _text(document_cfg.get("title_suffix")).strip() or "软件设计与使用说明书"
    cover = _as_mapping(document_cfg.get("cover", {}), "document.cover")
    for _ in range(_int(cover.get("top_blank_lines"), 5, "cover.top_blank_lines")):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(name)
    set_run_font(
        title_run,
        east_asia=style_info["body_font"],
        latin=style_info["latin_font"],
        size=_float(cover.get("software_name_size_pt"), 28.0, "cover.software_name_size_pt"),
        bold=bool(cover.get("software_name_bold", False)),
    )

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_before = Pt(18)
    subtitle_run = subtitle_paragraph.add_run(_text(document_cfg.get("cover_title")).strip() or title_suffix)
    set_run_font(
        subtitle_run,
        east_asia=style_info["body_font"],
        latin=style_info["latin_font"],
        size=_float(cover.get("document_title_size_pt"), 22.0, "cover.document_title_size_pt"),
        bold=bool(cover.get("document_title_bold", False)),
    )

    for _ in range(_int(cover.get("middle_blank_lines"), 3, "cover.middle_blank_lines")):
        document.add_paragraph()

    version_paragraph = document.add_paragraph()
    version_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_paragraph.paragraph_format.space_before = Pt(40)
    version_run = version_paragraph.add_run(version)
    set_run_font(
        version_run,
        east_asia=style_info["body_font"],
        latin=style_info["latin_font"],
        size=_float(cover.get("version_size_pt"), 24.0, "cover.version_size_pt"),
        bold=bool(cover.get("version_bold", False)),
    )

    owner = _text(software.get("owner") or cover.get("owner")).strip()
    date = _text(software.get("document_date") or cover.get("date")).strip()
    for value in (owner, date):
        if value:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            set_run_font(
                run,
                east_asia=style_info["body_font"],
                latin=style_info["latin_font"],
                size=_float(cover.get("auxiliary_size_pt"), 12.0, "cover.auxiliary_size_pt"),
            )
    document.add_page_break()


def add_toc(document: Document, document_cfg: dict[str, Any], style_info: dict[str, Any]) -> None:
    toc = _as_mapping(document_cfg.get("toc", {}), "document.toc")
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(_float(toc.get("title_after_pt"), 18.0, "toc.title_after_pt"))
    title_run = title_paragraph.add_run(_text(toc.get("title")).strip() or "目录")
    set_run_font(
        title_run,
        east_asia=style_info["body_font"],
        latin=style_info["latin_font"],
        size=_float(toc.get("title_size_pt"), 18.0, "toc.title_size_pt"),
    )
    min_level = _int(toc.get("min_level"), 1, "toc.min_level")
    max_level = _int(toc.get("max_level"), 2, "toc.max_level")
    if min_level < 1 or max_level < min_level or max_level > 9:
        raise ManualConfigError("toc levels must satisfy 1 <= min_level <= max_level <= 9")
    toc_paragraph = document.add_paragraph()
    field_run = add_field(
        toc_paragraph,
        f' TOC \\o "{min_level}-{max_level}" \\h \\z \\u ',
        _text(toc.get("placeholder")).strip() or "目录将在打开文档时更新",
    )
    set_run_font(
        field_run,
        east_asia=style_info["body_font"],
        latin=style_info["latin_font"],
        size=style_info["body_size"],
    )


def configure_body_footer(section: Any, document_cfg: dict[str, Any], style_info: dict[str, Any]) -> None:
    footer_cfg = _as_mapping(document_cfg.get("footer", {}), "document.footer")
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = _text(footer_cfg.get("prefix"))
    suffix = _text(footer_cfg.get("suffix"))
    if prefix:
        run = paragraph.add_run(prefix)
        set_run_font(run, east_asia=style_info["body_font"], latin=style_info["latin_font"], size=12)
    page_run = add_field(paragraph, " PAGE ", "1")
    set_run_font(page_run, east_asia=style_info["body_font"], latin=style_info["latin_font"], size=12)
    if bool(footer_cfg.get("show_section_total", False)):
        separator = paragraph.add_run(_text(footer_cfg.get("separator")) or " / ")
        set_run_font(separator, east_asia=style_info["body_font"], latin=style_info["latin_font"], size=12)
        total_run = add_field(paragraph, " SECTIONPAGES ", "1")
        set_run_font(total_run, east_asia=style_info["body_font"], latin=style_info["latin_font"], size=12)
    if suffix:
        run = paragraph.add_run(suffix)
        set_run_font(run, east_asia=style_info["body_font"], latin=style_info["latin_font"], size=12)


def build_document(payload: dict[str, Any], *, base_dir: Path) -> tuple[Document, RenderState]:
    software = _as_mapping(payload.get("software"), "software")
    name = _text(software.get("name")).strip()
    version = _text(software.get("version")).strip()
    if not name:
        raise ManualConfigError("software.name is required")
    if not version:
        raise ManualConfigError("software.version is required")
    document_cfg = _as_mapping(payload.get("document", {}), "document")
    layout = _as_mapping(document_cfg.get("layout", {}), "document.layout")
    style_cfg = _as_mapping(document_cfg.get("styles", {}), "document.styles")

    document = Document()
    configure_section(document.sections[0], layout)
    style_info = prepare_styles(document, style_cfg)
    enable_field_updates(document)
    clean_core_properties(document, software, document_cfg)
    add_cover(document, software, document_cfg, style_info)
    add_toc(document, document_cfg, style_info)

    body_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(body_section, layout)
    body_section.header.is_linked_to_previous = False
    restart_page_numbering(body_section, _int(document_cfg.get("body_page_start"), 1, "document.body_page_start"))
    configure_body_footer(body_section, document_cfg, style_info)
    max_width_cm = (
        body_section.page_width.cm - body_section.left_margin.cm - body_section.right_margin.cm
    )

    state = RenderState()
    top_blocks = _as_list(payload.get("blocks"), "blocks")
    chapters = _as_list(payload.get("chapters"), "chapters")
    if not top_blocks and not chapters:
        raise ManualConfigError("at least one top-level block or chapter is required")
    render_blocks(
        document,
        top_blocks,
        base_dir=base_dir,
        max_width_cm=max_width_cm,
        document_cfg=document_cfg,
        style_info=style_info,
        state=state,
    )
    for chapter in chapters:
        render_container(
            document,
            _as_mapping(chapter, "chapter"),
            heading_level=1,
            base_dir=base_dir,
            max_width_cm=max_width_cm,
            document_cfg=document_cfg,
            style_info=style_info,
            state=state,
        )
    return document, state


def load_payload(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8-sig"))
    except FileNotFoundError:
        raise
    except json.JSONDecodeError as exc:
        raise ManualConfigError(f"invalid JSON in {path}: {exc}") from exc
    return _as_mapping(data, "root")


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    example = r'''
JSON example:
  {
    "software": {"name": "Example System", "version": "V1.0"},
    "document": {"title_suffix": "Software Design and User Manual", "title_separator": " "},
    "chapters": [{
      "title": "1 Introduction",
      "sections": [{
        "title": "1.1 Overview",
        "blocks": [
          {"type": "paragraph", "text": "A substantial verified description."},
          {"type": "figure", "path": "figures/architecture.png",
           "caption": "System architecture", "width_cm": 15.0},
          {"type": "table", "caption": "Module summary",
           "headers": ["Module", "Purpose"], "rows": [["Core", "Processing"]]}
        ]
      }]
    }]
  }

Image paths are resolved relative to the JSON file unless --base-dir is set.
The generated TOC and page fields are dynamic and should be refreshed by
Word, WPS Office, or LibreOffice before final submission.
'''
    parser = argparse.ArgumentParser(
        description="Build an A4 software design and user manual DOCX from JSON.",
        epilog=example,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--input", required=True, type=Path, help="UTF-8 JSON manual specification")
    parser.add_argument("--output", required=True, type=Path, help="destination .docx file")
    parser.add_argument(
        "--base-dir",
        type=Path,
        help="base directory for relative figure paths (default: JSON file directory)",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    input_path = args.input.expanduser().resolve()
    output_path = args.output.expanduser().resolve()
    if output_path.suffix.lower() != ".docx":
        raise ManualConfigError("--output must use the .docx extension")
    base_dir = args.base_dir.expanduser().resolve() if args.base_dir else input_path.parent
    payload = load_payload(input_path)
    document, state = build_document(payload, base_dir=base_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    summary = {
        "output": str(output_path),
        "headings": state.headings,
        "paragraphs": state.paragraphs,
        "figures": state.figures,
        "tables": state.tables,
        "missing_images": state.missing_images,
    }
    print(json.dumps(summary, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (ManualConfigError, FileNotFoundError) as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise SystemExit(2)
