#!/usr/bin/env python3
"""Audit the three software-copyright DOCX artifacts against one fact source."""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
from datetime import date
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from lxml import etree


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
EP = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
DC = "http://purl.org/dc/elements/1.1/"
CP = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
PLACEHOLDER = re.compile(r"TODO|FIXME|TBD|待插入|待补|后期补图|占位符|示例系统", re.I)
MOJIBAKE = re.compile(
    r"[�\ue000-\uf8ff]|(?:[鏈鐢璁鍖缃锛銆搴櫌涓鍦鐨鏄埛綍嶅瓨ㄣ].{0,3}){3,}"
)
HEX_SHA256 = re.compile(r"[0-9a-f]{64}", re.I)
XML_PARSER = etree.XMLParser(resolve_entities=False, no_network=True, huge_tree=False)
TEXT_PARTS = re.compile(r"^word/(?:document|header\d+|footer\d+|footnotes|endnotes)\.xml$")
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
DEFAULT_REQUIRED_MANUAL_CHAPTERS = [
    "软件设计说明",
    "软件使用说明",
    "软硬件运行环境",
    "知识产权声明",
]
FIGURE_REVIEW_FLAGS = {
    "software_name_correct",
    "old_name_absent",
    "pii_absent_or_redacted",
    "credential_absent",
    "error_absent",
    "splice_absent",
    "manually_reviewed",
}


def load_json(path: Path | None) -> dict[str, Any]:
    if path is None:
        return {}
    value = json.loads(path.read_text(encoding="utf-8-sig"))
    if not isinstance(value, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return value


def norm_version(value: str) -> str:
    return str(value or "").strip().lstrip("vV")


def norm_date(value: str) -> str:
    compact = str(value or "").strip().replace("年", "-").replace("月", "-").replace("日", "")
    compact = compact.replace("/", "-").replace(".", "-")
    try:
        return date.fromisoformat(compact).isoformat()
    except ValueError:
        return ""


def as_int(value: Any, default: int = -1) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def parse_xml(data: bytes) -> etree._Element:
    return etree.fromstring(data, parser=XML_PARSER)


def field_cache_values(root: etree._Element, field_name: str) -> list[str]:
    """Return cached values for complex and simple Word fields."""
    results: list[str] = []
    matcher = re.compile(rf"(?<![A-Z]){re.escape(field_name.upper())}(?![A-Z])")
    for paragraph in root.iter(f"{{{W}}}p"):
        stack: list[dict[str, Any]] = []
        for node in paragraph.iter():
            if node.tag == f"{{{W}}}fldChar":
                kind = node.get(f"{{{W}}}fldCharType")
                if kind == "begin":
                    stack.append({"instruction": [], "separated": False, "match": False, "value": []})
                elif kind == "separate" and stack:
                    stack[-1]["separated"] = True
                    stack[-1]["match"] = bool(matcher.search("".join(stack[-1]["instruction"]).upper()))
                elif kind == "end" and stack:
                    context = stack.pop()
                    if context["match"]:
                        results.append("".join(context["value"]).strip())
                continue
            if node.tag == f"{{{W}}}instrText" and stack and not stack[-1]["separated"]:
                stack[-1]["instruction"].append(node.text or "")
            elif node.tag == f"{{{W}}}t" and stack:
                for context in stack:
                    if context["separated"] and context["match"]:
                        context["value"].append(node.text or "")
    for simple in root.iter(f"{{{W}}}fldSimple"):
        if matcher.search(simple.get(f"{{{W}}}instr", "").upper()):
            results.append("".join(simple.itertext()).strip())
    return results


def docx_package_info(path: Path) -> dict[str, Any]:
    with ZipFile(path) as package:
        if sum(info.file_size for info in package.infolist()) > 1_000_000_000:
            raise ValueError(f"DOCX expands beyond the 1 GB audit limit: {path}")
        bad = package.testzip()
        if bad:
            raise ValueError(f"Corrupt DOCX entry {bad}: {path}")
        required = {"docProps/app.xml", "docProps/core.xml", "word/document.xml", "word/settings.xml"}
        missing = sorted(required - set(package.namelist()))
        if missing:
            raise ValueError(f"DOCX is missing required parts {missing}: {path}")
        app = parse_xml(package.read("docProps/app.xml"))
        core = parse_xml(package.read("docProps/core.xml"))
        settings = parse_xml(package.read("word/settings.xml"))
        document = parse_xml(package.read("word/document.xml"))
        body_text: list[str] = []
        auxiliary_text: list[str] = []
        metadata_text: list[str] = []
        instructions: list[str] = []
        numpages_values: list[str] = []
        text_colors: list[dict[str, str]] = []
        revision_nodes = 0
        hidden_nodes = 0
        for name in package.namelist():
            if TEXT_PARTS.match(name):
                root = document if name == "word/document.xml" else parse_xml(package.read(name))
                for node in root.xpath(".//w:rPr/w:color", namespaces={"w": W}):
                    text_colors.append(
                        {
                            "part": name,
                            "value": node.get(f"{{{W}}}val", ""),
                            "theme": node.get(f"{{{W}}}themeColor", ""),
                        }
                    )
                instructions.extend(" ".join(item.split()) for item in root.xpath(".//w:instrText/text()", namespaces={"w": W}))
                instructions.extend(" ".join(item.split()) for item in root.xpath(".//w:fldSimple/@w:instr", namespaces={"w": W}))
                values = [str(item) for item in root.xpath(".//w:t/text()", namespaces={"w": W})]
                (body_text if name == "word/document.xml" else auxiliary_text).extend(values)
                numpages_values.extend(field_cache_values(root, "NUMPAGES"))
                revision_nodes += len(root.findall(f".//{{{W}}}ins")) + len(root.findall(f".//{{{W}}}del"))
                hidden_nodes += len(root.findall(f".//{{{W}}}vanish"))
            elif (name.startswith("docProps/") or name.startswith("customXml/")) and name.endswith(".xml"):
                try:
                    metadata_text.extend(str(item) for item in parse_xml(package.read(name)).itertext())
                except etree.XMLSyntaxError:
                    metadata_text.append("<invalid-xml>")
        rel_errors: list[dict[str, str]] = []
        for name in package.namelist():
            if not name.endswith(".rels"):
                continue
            root = parse_xml(package.read(name))
            for relationship in root:
                if relationship.get("TargetMode") == "External":
                    # Never copy a possibly credential-bearing URL into the report.
                    rel_errors.append({"part": name, "type": relationship.get("Type", "")})
        update_fields = settings.find(f"{{{W}}}updateFields")
        return {
            "pages": int(app.findtext(f"{{{EP}}}Pages") or 0),
            "creator": core.findtext(f"{{{DC}}}creator") or "",
            "last_modified_by": core.findtext(f"{{{CP}}}lastModifiedBy") or "",
            "title": core.findtext(f"{{{DC}}}title") or "",
            "text": "\n".join(body_text),
            "header_footer_text": "\n".join(auxiliary_text),
            "metadata_text": "\n".join(metadata_text),
            "field_instructions": instructions,
            "numpages_values": numpages_values,
            "text_colors": text_colors,
            "update_fields": update_fields is not None
            and update_fields.get(f"{{{W}}}val", "true").lower() not in {"0", "false", "off"},
            "media_count": len([name for name in package.namelist() if name.startswith("word/media/") and not name.endswith("/")]),
            "page_breaks": len(document.findall(f".//{{{W}}}br[@{{{W}}}type='page']")),
            "comments_parts": [name for name in package.namelist() if "comment" in name.lower()],
            "revision_nodes": revision_nodes,
            "hidden_nodes": hidden_nodes,
            "external_relationships": rel_errors,
        }


def add_issue(container: list[dict[str, Any]], code: str, message: str, **details: Any) -> None:
    item: dict[str, Any] = {"code": code, "message": message}
    item.update(details)
    container.append(item)


def normalized_windows(lines: list[str], width: int) -> dict[str, list[int]]:
    normalized = [re.sub(r"\s+", " ", line.strip()) for line in lines]
    windows: dict[str, list[int]] = {}
    for index in range(len(normalized) - width + 1):
        digest = hashlib.sha256("\n".join(normalized[index : index + width]).encode("utf-8")).hexdigest()
        windows.setdefault(digest, []).append(index + 1)
    return windows


def pdf_page_count(path: Path) -> int | None:
    if not path.exists():
        return None
    try:
        import fitz

        with fitz.open(path) as document:
            return len(document)
    except Exception:
        return None


def decode_source(raw: bytes) -> str:
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    if raw.startswith((b"\xff\xfe", b"\xfe\xff")):
        return raw.decode("utf-16")
    for encoding in ("utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("source", raw, 0, min(1, len(raw)), "unsupported source encoding")


def current_policy_files(root: Path, policy: dict[str, Any]) -> set[str]:
    extensions = {str(value).lower() for value in policy.get("include_extensions", [])}
    excluded_parts = {str(value).lower() for value in policy.get("excluded_parts", [])}
    excluded_suffixes = tuple(str(value).lower() for value in policy.get("excluded_suffixes", []))
    exclude_globs = [str(value) for value in policy.get("exclude_globs", [])]
    include_globs = [str(value) for value in policy.get("include_globs", [])]
    result: set[str] = set()
    for path in root.rglob("*"):
        if not path.is_file() or (extensions and path.suffix.lower() not in extensions):
            continue
        relative = path.relative_to(root)
        if {part.lower() for part in relative.parts[:-1]} & excluded_parts:
            continue
        if excluded_suffixes and path.name.lower().endswith(excluded_suffixes):
            continue
        if any(relative.match(pattern) for pattern in exclude_globs):
            continue
        if include_globs and not any(relative.match(pattern) for pattern in include_globs):
            continue
        result.add(relative.as_posix())
    return result


def verify_snapshot(manifest: dict[str, Any]) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    errors: list[dict[str, Any]] = []
    stats: dict[str, Any] = {"files": 0, "program_nonblank_lines": 0, "snapshot_sha256": ""}
    root_value = manifest.get("project_root")
    if not isinstance(root_value, str) or not root_value.strip():
        add_issue(errors, "manifest-root", "Frozen source manifest has no project_root")
        return errors, stats
    root = Path(root_value).expanduser().resolve()
    if not root.is_dir():
        add_issue(errors, "manifest-root", "Frozen project root no longer exists", root=str(root))
        return errors, stats
    entries = manifest.get("files")
    if not isinstance(entries, list) or not entries:
        add_issue(errors, "manifest-files", "Frozen source manifest contains no source files")
        return errors, stats
    seen: set[str] = set()
    snapshot = hashlib.sha256()
    for position, item in enumerate(entries):
        if not isinstance(item, dict) or not isinstance(item.get("path"), str):
            add_issue(errors, "manifest-entry", "Malformed source manifest entry", index=position)
            continue
        relative = item["path"].replace("\\", "/")
        if relative in seen:
            add_issue(errors, "manifest-duplicate", "Source manifest contains a duplicate path", file=relative)
            continue
        seen.add(relative)
        unresolved = root / relative
        candidate = unresolved.resolve()
        try:
            candidate.relative_to(root)
        except ValueError:
            add_issue(errors, "manifest-path-escape", "Source manifest path escapes project_root", file=relative)
            continue
        relative_parts = Path(relative).parts
        if any((root.joinpath(*relative_parts[:index])).is_symlink() for index in range(1, len(relative_parts) + 1)):
            add_issue(errors, "source-symlink", "Frozen source must not be a symbolic link", file=relative)
        if not candidate.is_file():
            add_issue(errors, "source-missing", "A frozen source file is missing", file=relative)
            continue
        raw = candidate.read_bytes()
        digest = hashlib.sha256(raw).hexdigest()
        if not HEX_SHA256.fullmatch(str(item.get("sha256", ""))) or digest.lower() != str(item.get("sha256", "")).lower():
            add_issue(errors, "source-changed", "Source changed after the manifest was frozen", file=relative)
        try:
            text = decode_source(raw)
        except UnicodeDecodeError:
            add_issue(errors, "source-decode", "Frozen source is no longer strictly decodable", file=relative)
            continue
        nonblank = sum(1 for line in text.splitlines() if line.strip())
        if nonblank != as_int(item.get("nonblank_lines")):
            add_issue(errors, "source-line-change", "Source nonblank line count differs from the manifest", file=relative)
        if len(text.splitlines()) != as_int(item.get("physical_lines")):
            add_issue(errors, "source-physical-change", "Source physical line count differs from the manifest", file=relative)
        snapshot.update(f"{relative}\0{digest}\0{nonblank}\n".encode("utf-8"))
        stats["program_nonblank_lines"] += nonblank
        stats["files"] += 1
    stats["snapshot_sha256"] = snapshot.hexdigest()
    if stats["program_nonblank_lines"] != as_int(manifest.get("program_nonblank_lines")):
        add_issue(errors, "manifest-line-total", "Manifest total source line count is inconsistent")
    if stats["snapshot_sha256"].lower() != str(manifest.get("snapshot_sha256", "")).lower():
        add_issue(errors, "manifest-snapshot", "Manifest snapshot SHA-256 is inconsistent")
    policy = manifest.get("source_policy")
    if not isinstance(policy, dict) or not policy.get("include_extensions"):
        add_issue(errors, "manifest-policy", "Manifest lacks the discovery policy needed to detect added source files")
    else:
        current = current_policy_files(root, policy)
        if current != seen:
            add_issue(
                errors,
                "source-set-changed",
                "Source files were added to or removed from the frozen source set",
                added=sorted(current - seen),
                removed=sorted(seen - current),
            )
    return errors, stats


def legal_missing(config: dict[str, Any]) -> list[str]:
    legal = config.get("legal") or {}
    applicant = legal.get("applicant") or config.get("applicant") or {}
    holders = legal.get("rights_holders") or config.get("copyright_owners") or []
    missing: list[str] = []
    if not isinstance(holders, list) or not holders:
        missing.append("copyright_owners")
    else:
        for index, holder in enumerate(holders):
            if not isinstance(holder, dict):
                missing.append(f"copyright_owners[{index}]")
                continue
            for key in ("name", "credential_no", "address", "nationality"):
                if not holder.get(key):
                    missing.append(f"copyright_owners[{index}].{key}")
    for key in ("name", "address", "postal_code", "contact", "mobile", "email"):
        if not applicant.get(key):
            missing.append(f"applicant.{key}")
    return missing


def evidence_features(data: dict[str, Any]) -> list[dict[str, Any]]:
    raw = data.get("features", data.get("feature_evidence", []))
    if isinstance(raw, dict):
        return [dict(value, feature_id=str(key)) for key, value in raw.items() if isinstance(value, dict)]
    return [value for value in raw if isinstance(value, dict)] if isinstance(raw, list) else []


def canonical_layer(value: str) -> str:
    clean = re.sub(r"[^a-z]", "", value.lower())
    if clean in {"ui", "interaction", "frontend", "client", "view", "page", "screen", "component", "cli", "command", "desktop"}:
        return "ui"
    if clean in {"api", "route", "router", "controller", "endpoint", "handler"}:
        return "api"
    if clean in {"business", "service", "domain", "core", "workflow", "persistence", "repository", "database", "dao"}:
        return "domain"
    return ""


def feature_code_references(feature: dict[str, Any]) -> list[dict[str, Any]]:
    references: list[dict[str, Any]] = []
    grouped = feature.get("evidence")
    if isinstance(grouped, dict):
        for layer, values in grouped.items():
            if layer in {"runtime", "screenshots", "notes"}:
                continue
            items = values if isinstance(values, list) else [values]
            for value in items:
                if isinstance(value, str):
                    references.append({"path": value, "layer": layer})
                elif isinstance(value, dict):
                    references.append({**value, "layer": value.get("layer", layer)})
    raw = feature.get("code_references", feature.get("code_refs", []))
    if isinstance(raw, list):
        for value in raw:
            if isinstance(value, str):
                references.append({"path": value, "layer": ""})
            elif isinstance(value, dict):
                references.append(dict(value))
    return references


def extract_manual_feature_refs(value: Any) -> tuple[set[str], set[str]]:
    claimed: set[str] = set()
    core: set[str] = set()
    if isinstance(value, dict):
        for key, child in value.items():
            if key in {"feature_id", "evidence_feature_id"} and isinstance(child, str) and child.strip():
                claimed.add(child.strip())
            elif key in {"feature_ids", "evidence_feature_ids"} and isinstance(child, list):
                claimed.update(str(item).strip() for item in child if str(item).strip())
            elif key == "core_feature_id" and isinstance(child, str) and child.strip():
                core.add(child.strip())
            elif key == "core_feature_ids" and isinstance(child, list):
                core.update(str(item).strip() for item in child if str(item).strip())
            extract_claimed, extract_core = extract_manual_feature_refs(child)
            claimed.update(extract_claimed)
            core.update(extract_core)
    elif isinstance(value, list):
        for child in value:
            extract_claimed, extract_core = extract_manual_feature_refs(child)
            claimed.update(extract_claimed)
            core.update(extract_core)
    claimed.update(core)
    return claimed, core


def screenshot_paths(feature: dict[str, Any], base: Path) -> list[Path]:
    raw = feature.get("screenshots")
    if raw is None and isinstance(feature.get("evidence"), dict):
        raw = feature["evidence"].get("screenshots")
    if raw is None:
        raw = feature.get("screenshot")
    items = raw if isinstance(raw, list) else ([raw] if raw else [])
    result: list[Path] = []
    for item in items:
        if isinstance(item, str):
            value = item
            kind = "screenshot"
        elif isinstance(item, dict):
            value = item.get("path") or item.get("file")
            kind = str(item.get("kind") or item.get("type") or "screenshot").lower()
        else:
            continue
        if not isinstance(value, str) or kind in {"diagram", "generated", "mockup", "placeholder"}:
            continue
        path = Path(value).expanduser()
        result.append((path if path.is_absolute() else base / path).resolve())
    return result


def valid_screenshot(path: Path) -> bool:
    if not path.is_file() or path.stat().st_size <= 0 or path.suffix.lower() not in IMAGE_SUFFIXES:
        return False
    try:
        from PIL import Image

        with Image.open(path) as image:
            width, height = image.size
            image.verify()
        return width >= 100 and height >= 100
    except Exception:
        return False


def normalized_chapter_title(value: Any) -> str:
    return re.sub(
        r"^\s*(?:第\s*)?[0-9一二三四五六七八九十]+\s*(?:章)?[\s.、．-]*",
        "",
        str(value or "").strip(),
    )


def audit_manual_structure(
    manual_content: dict[str, Any], manual_cfg: dict[str, Any]
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    errors: list[dict[str, Any]] = []
    required_titles = [normalized_chapter_title(value) for value in DEFAULT_REQUIRED_MANUAL_CHAPTERS]
    raw_required = manual_cfg.get("required_chapters", DEFAULT_REQUIRED_MANUAL_CHAPTERS)
    configured_titles = (
        [normalized_chapter_title(value) for value in raw_required if str(value or "").strip()]
        if isinstance(raw_required, list)
        else []
    )
    raw_chapters = manual_content.get("chapters")
    chapters = raw_chapters if isinstance(raw_chapters, list) else []
    actual_titles = [
        normalized_chapter_title(chapter.get("title"))
        for chapter in chapters
        if isinstance(chapter, dict)
    ]
    if configured_titles != required_titles:
        add_issue(errors, "manual-required-chapters", "The fixed four-chapter manual structure cannot be overridden", configured=configured_titles)
    if actual_titles != required_titles:
        add_issue(
            errors,
            "manual-chapters",
            "Manual chapters must use the complete software-copyright structure and end with the environment and IP chapters",
            expected=required_titles,
            actual=actual_titles,
        )
    return errors, {"required_chapters": required_titles, "actual_chapters": actual_titles}


def audit_figures_manifest(
    figures: dict[str, Any], figures_base: Path, evidence: dict[str, Any]
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    errors: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    raw_figures = figures.get("figures")
    items = raw_figures if isinstance(raw_figures, list) else []
    if not items:
        add_issue(errors, "figures-empty", "figures_manifest.json contains no figures")
    evidence_index = {
        str(feature.get("feature_id") or feature.get("id") or "").strip(): feature
        for feature in evidence_features(evidence)
        if str(feature.get("feature_id") or feature.get("id") or "").strip()
    }
    runtime_coverage: set[str] = set()
    identifiers: set[str] = set()
    runtime_count = 0
    for item in items:
        if not isinstance(item, dict):
            add_issue(errors, "figure-entry", "Figure manifest entry must be an object")
            continue
        figure_id = str(item.get("id") or "").strip()
        if not figure_id or figure_id in identifiers:
            add_issue(errors, "figure-id", "Figure id is missing or duplicated", figure_id=figure_id)
        identifiers.add(figure_id)
        kind = str(item.get("kind") or "").strip().lower()
        if kind == "placeholder":
            add_issue(errors, "figure-placeholder", "Placeholder figure remains in the final figure manifest", figure_id=figure_id)
            continue
        raw_path = item.get("path")
        if not isinstance(raw_path, str) or not raw_path.strip():
            add_issue(errors, "figure-path", "Figure path is missing", figure_id=figure_id)
            continue
        path = Path(raw_path).expanduser()
        path = (path if path.is_absolute() else figures_base / path).resolve()
        if not path.is_file() or path.stat().st_size <= 0:
            add_issue(errors, "figure-missing", "Figure file does not exist", figure_id=figure_id, path=str(path))
            continue
        if kind == "runtime_screenshot":
            runtime_count += 1
            if not valid_screenshot(path):
                add_issue(errors, "figure-invalid-screenshot", "Runtime screenshot is unreadable or too small", figure_id=figure_id)
            feature_ids = {
                str(value).strip()
                for value in item.get("feature_ids", [])
                if str(value).strip()
            } if isinstance(item.get("feature_ids"), list) else set()
            if not feature_ids:
                add_issue(errors, "figure-feature-links", "Runtime screenshot has no feature_ids", figure_id=figure_id)
            for feature_id in feature_ids:
                if feature_id not in evidence_index:
                    add_issue(errors, "figure-feature-unknown", "Screenshot references an unknown feature", figure_id=figure_id, feature_id=feature_id)
                else:
                    runtime_coverage.add(feature_id)
            capture_tool = str(item.get("capture_tool") or item.get("capture_method") or "").strip()
            if not capture_tool:
                add_issue(errors, "figure-capture-provenance", "Runtime screenshot does not record its browser/capture tool", figure_id=figure_id)
            review = item.get("review") if isinstance(item.get("review"), dict) else {}
            missing_flags = sorted(flag for flag in FIGURE_REVIEW_FLAGS if review.get(flag) is not True)
            if missing_flags:
                add_issue(errors, "figure-review", "Runtime screenshot review is incomplete", figure_id=figure_id, fields=missing_flags)
        elif kind not in {"generated_diagram", "diagram"}:
            add_issue(warnings, "figure-kind", "Figure kind is not a recognized runtime screenshot or generated diagram", figure_id=figure_id, kind=kind)
    required_runtime = {
        feature_id
        for feature_id, feature in evidence_index.items()
        if as_int(feature.get("evidence_level", feature.get("level")), -1) == 4
        and feature.get("include_in_manual", True) is not False
        and str(feature.get("status") or "implemented").strip().lower()
        not in {"excluded", "not_implemented", "unimplemented"}
    }
    missing_runtime = sorted(required_runtime - runtime_coverage)
    if missing_runtime:
        add_issue(errors, "figure-runtime-coverage", "Level-4 manual features lack reviewed runtime screenshots", feature_ids=missing_runtime)
    return errors, warnings, {
        "figures": len(items),
        "runtime_screenshots": runtime_count,
        "runtime_feature_ids": sorted(runtime_coverage),
    }


def audit_feature_evidence(
    evidence: dict[str, Any],
    evidence_base: Path,
    manual_content: dict[str, Any],
    manifest: dict[str, Any],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    errors: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    features = evidence_features(evidence)
    indexed: dict[str, dict[str, Any]] = {}
    source_entries = {
        str(item.get("path", "")).replace("\\", "/"): item
        for item in manifest.get("files", [])
        if isinstance(item, dict)
    }
    for feature in features:
        feature_id = str(feature.get("feature_id") or feature.get("id") or "").strip()
        if not feature_id:
            add_issue(errors, "evidence-feature-id", "Evidence feature has no feature_id")
            continue
        if feature_id in indexed:
            add_issue(errors, "evidence-feature-duplicate", "Duplicate feature_id in evidence", feature_id=feature_id)
            continue
        indexed[feature_id] = feature
        try:
            level = int(feature.get("evidence_level", feature.get("level", -1)))
        except (TypeError, ValueError):
            level = -1
        if level not in range(5):
            add_issue(errors, "evidence-level", "Evidence level must be an integer from 0 through 4", feature_id=feature_id)
            continue
        references = feature_code_references(feature)
        layers: set[str] = set()
        for reference in references:
            layer = canonical_layer(str(reference.get("layer") or reference.get("kind") or reference.get("type") or ""))
            if layer:
                layers.add(layer)
            value = reference.get("path") or reference.get("file")
            if not isinstance(value, str) or not value.strip():
                add_issue(errors, "evidence-code-path", "Code evidence reference has no source path", feature_id=feature_id)
                continue
            relative = value.replace("\\", "/")
            while relative.startswith("./"):
                relative = relative[2:]
            if relative.startswith("/") or ".." in Path(relative).parts:
                add_issue(errors, "evidence-code-path", "Code evidence path must remain inside the frozen project", feature_id=feature_id)
                continue
            entry = source_entries.get(relative)
            if entry is None:
                add_issue(errors, "evidence-code-unfrozen", "Code evidence is not in the frozen source manifest", feature_id=feature_id, file=relative)
                continue
            line_value = reference.get("line", reference.get("start_line"))
            if line_value is not None:
                try:
                    line = int(line_value)
                    if line <= 0 or line > int(entry.get("physical_lines", 0)):
                        raise ValueError
                except (TypeError, ValueError):
                    add_issue(errors, "evidence-code-line", "Code evidence line is outside the frozen source file", feature_id=feature_id, file=relative)
        raw_required = feature.get("required_layers", ["ui", "api", "domain"])
        if not isinstance(raw_required, list) or not raw_required:
            add_issue(errors, "evidence-required-layers", "required_layers must be a non-empty array", feature_id=feature_id)
            required_layers = {"ui", "api", "domain"}
        else:
            required_layers = {canonical_layer(str(value)) for value in raw_required}
            if "" in required_layers:
                add_issue(errors, "evidence-required-layers", "required_layers contains an unknown layer", feature_id=feature_id)
                required_layers.discard("")
        if level >= 3 and not required_layers.issubset(layers):
            add_issue(
                errors,
                "evidence-chain",
                "Level 3/4 feature does not satisfy its required interaction/API/domain source chain",
                feature_id=feature_id,
                layers=sorted(layers),
                required_layers=sorted(required_layers),
            )
        if level == 4:
            runtime_verified = feature.get("runtime_verified") is True or (
                isinstance(feature.get("runtime"), dict) and feature["runtime"].get("verified") is True
            )
            if not runtime_verified:
                add_issue(errors, "evidence-runtime", "Level 4 feature lacks runtime_verified=true", feature_id=feature_id)
            screenshots = screenshot_paths(feature, evidence_base)
            valid = [path for path in screenshots if valid_screenshot(path)]
            if not valid:
                add_issue(errors, "evidence-screenshot", "Level 4 feature lacks an existing real screenshot file", feature_id=feature_id)
    claimed, core = extract_manual_feature_refs(manual_content)
    if not claimed:
        add_issue(errors, "manual-evidence-links", "manual_content does not reference any feature_id")
    for feature_id in sorted(claimed):
        feature = indexed.get(feature_id)
        if feature is None:
            add_issue(errors, "manual-feature-unknown", "Manual references a feature absent from feature evidence", feature_id=feature_id)
            continue
        try:
            level = int(feature.get("evidence_level", feature.get("level", -1)))
        except (TypeError, ValueError):
            level = -1
        if level < 3:
            add_issue(errors, "manual-feature-unclaimable", "Manual claims a feature below evidence level 3", feature_id=feature_id, level=level)
        if feature_id in core and level < 4:
            add_issue(errors, "manual-core-feature", "Core/screenshot manual feature must have evidence level 4", feature_id=feature_id, level=level)
    required_manual: set[str] = set()
    deliberately_excluded: set[str] = set()
    for feature_id, feature in indexed.items():
        level = as_int(feature.get("evidence_level", feature.get("level")), -1)
        status = str(feature.get("status") or "implemented").strip().lower()
        if feature.get("include_in_manual") is False:
            deliberately_excluded.add(feature_id)
            reason = str(feature.get("manual_exclusion_reason") or feature.get("exclusion_reason") or "").strip()
            if level >= 3 and not reason:
                add_issue(errors, "manual-feature-exclusion", "Implemented feature excluded from the manual without a reason", feature_id=feature_id)
            continue
        if level >= 3 and status not in {"excluded", "not_implemented", "unimplemented"}:
            required_manual.add(feature_id)
    missing_claimable = sorted(required_manual - claimed)
    if missing_claimable:
        add_issue(errors, "manual-feature-coverage", "Implemented user functions are missing from manual_content", feature_ids=missing_claimable)
    remaining_unreferenced = sorted(set(indexed) - claimed - deliberately_excluded - required_manual)
    if remaining_unreferenced:
        add_issue(warnings, "evidence-unreferenced", "Low-evidence or non-implemented features are not referenced by manual_content", feature_ids=remaining_unreferenced)
    return errors, warnings, {"features": len(indexed), "manual_feature_ids": sorted(claimed), "core_feature_ids": sorted(core)}


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit manual, application draft and source-code deposit as one consistent submission kit.")
    parser.add_argument("--config", required=True, type=Path)
    parser.add_argument("--facts", type=Path)
    parser.add_argument("--manifest", type=Path)
    parser.add_argument("--deposit-manifest", type=Path)
    parser.add_argument("--evidence", type=Path, help="feature_evidence.json")
    parser.add_argument("--figures", type=Path, help="figures_manifest.json with screenshot provenance and review")
    parser.add_argument("--manual-content", type=Path, help="manual_content.json with feature_id links")
    parser.add_argument("--manual", required=True, type=Path)
    parser.add_argument("--application", required=True, type=Path)
    parser.add_argument("--code", required=True, type=Path)
    parser.add_argument("--pdf-dir", type=Path)
    parser.add_argument("--forbidden-term", action="append", default=[])
    parser.add_argument("--report-json", required=True, type=Path)
    parser.add_argument("--report-md", type=Path)
    args = parser.parse_args()

    config = load_json(args.config)
    facts = load_json(args.facts)
    manifest = load_json(args.manifest)
    deposit_manifest = load_json(args.deposit_manifest)
    evidence = load_json(args.evidence)
    figures = load_json(args.figures)
    manual_content = load_json(args.manual_content)
    software = config.get("software") or {}
    name = str(software.get("full_name") or "").strip()
    version = str(software.get("version") or "").strip()
    completion_date = str(software.get("completion_date") or "").strip()
    expected_program_lines = int(manifest.get("program_nonblank_lines") or facts.get("program_nonblank_lines") or software.get("source_line_count") or 0)
    confirmations = config.get("confirmations") or {}
    errors: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []

    for label, path in (
        ("facts", args.facts),
        ("source-manifest", args.manifest),
        ("deposit-manifest", args.deposit_manifest),
        ("feature-evidence", args.evidence),
        ("figures-manifest", args.figures),
        ("manual-content", args.manual_content),
    ):
        if path is None:
            add_issue(errors, "missing-audit-input", "Required reproducibility/evidence input was not supplied", input=label)

    if not name:
        add_issue(errors, "missing-name", "Registered full software name is missing")
    if not version:
        add_issue(errors, "missing-version", "Registered version is missing")
    if not completion_date:
        add_issue(errors, "missing-completion-date", "Completion date is missing")
    else:
        try:
            parsed_completion = date.fromisoformat(completion_date)
            if parsed_completion > date.today():
                add_issue(errors, "future-completion-date", "Completion date cannot be in the future", value=completion_date)
        except ValueError:
            add_issue(errors, "invalid-completion-date", "Completion date must use YYYY-MM-DD", value=completion_date)
    if not confirmations.get("completion_date_confirmed", False):
        add_issue(errors, "completion-date-unconfirmed", "Completion date is a legal fact and has not been confirmed")
    if not confirmations.get("rights_confirmed", False):
        add_issue(errors, "rights-unconfirmed", "Work nature, development mode and rights scope have not been confirmed")
    if not confirmations.get("screenshots_reviewed", False):
        add_issue(errors, "screenshots-unreviewed", "Screenshots require OCR or human review for old names, seams and sensitive data")
    if not confirmations.get("official_rules_verified", False):
        add_issue(errors, "official-rules-unverified", "Current receiving-platform rules and official form requirements have not been verified")
    if not confirmations.get("source_scope_confirmed", False):
        add_issue(errors, "source-scope-unconfirmed", "The first-party source inclusion/exclusion scope has not been explicitly confirmed")
    if not confirmations.get("software_identity_confirmed", False):
        add_issue(errors, "software-identity-unconfirmed", "Software name, version and completion identity have not been explicitly confirmed")
    for field in legal_missing(config):
        add_issue(errors, "missing-legal-field", "Required legal/application field is missing", field=field)

    for item in facts.get("errors", []):
        add_issue(errors, "project-fact-error", "Project scan reported an error", detail=item)
    for item in facts.get("risks", []):
        if item.get("severity") == "error":
            add_issue(errors, "project-source-risk", "Project scan reported a blocking source risk", detail=item)
        elif item.get("severity") == "warning":
            add_issue(warnings, "project-source-warning", "Project scan reported a source warning", detail=item)
    snapshot_stats: dict[str, Any] = {}
    if manifest:
        snapshot_errors, snapshot_stats = verify_snapshot(manifest)
        errors.extend(snapshot_errors)
        if facts.get("snapshot_sha256") and facts.get("snapshot_sha256") != manifest.get("snapshot_sha256"):
            add_issue(errors, "facts-manifest-snapshot", "Project facts and source manifest have different snapshot hashes")
        if facts.get("program_nonblank_lines") and int(facts["program_nonblank_lines"]) != int(manifest.get("program_nonblank_lines", -1)):
            add_issue(errors, "facts-manifest-lines", "Project facts and source manifest have different source line totals")
    evidence_stats: dict[str, Any] = {}
    if evidence and manual_content and manifest and args.evidence:
        evidence_errors, evidence_warnings, evidence_stats = audit_feature_evidence(
            evidence, args.evidence.parent.resolve(), manual_content, manifest
        )
        errors.extend(evidence_errors)
        warnings.extend(evidence_warnings)
    figure_stats: dict[str, Any] = {}
    if args.figures:
        figure_errors, figure_warnings, figure_stats = audit_figures_manifest(
            figures, args.figures.parent.resolve(), evidence
        )
        errors.extend(figure_errors)
        warnings.extend(figure_warnings)

    artifacts = {
        "manual": (args.manual, docx_package_info(args.manual)),
        "application": (args.application, docx_package_info(args.application)),
        "code": (args.code, docx_package_info(args.code)),
    }
    old_names = software.get("old_names") or software.get("forbidden_names") or []
    if isinstance(old_names, str):
        old_names = [old_names]
    forbidden_terms = list(dict.fromkeys([
        *args.forbidden_term,
        *((config.get("source") or {}).get("forbidden_terms") or []),
        *(str(value) for value in old_names),
    ]))
    for kind, (path, info) in artifacts.items():
        combined = f"{info['text']}\n{info['header_footer_text']}\n{info['metadata_text']}"
        if name and name not in combined:
            add_issue(errors, "name-missing-from-artifact", "Registered name not found in artifact", artifact=kind)
        if version and version not in combined:
            add_issue(errors, "version-missing-from-artifact", "Registered version not found in artifact", artifact=kind)
        if PLACEHOLDER.search(combined):
            add_issue(errors, "placeholder", "Placeholder text remains", artifact=kind)
        if MOJIBAKE.search(combined):
            add_issue(errors, "mojibake", "Possible mojibake remains", artifact=kind)
        for term in forbidden_terms:
            if term and term in combined:
                add_issue(errors, "forbidden-term", "Old or forbidden term remains", artifact=kind, term=term)
        if info["creator"] or info["last_modified_by"]:
            add_issue(errors, "unclean-metadata", "Creator/lastModifiedBy metadata is not blank", artifact=kind)
        if info["comments_parts"] or info["revision_nodes"] or info["hidden_nodes"]:
            add_issue(errors, "review-artifacts", "Comments, revisions or hidden text remain", artifact=kind)
        if info["external_relationships"]:
            add_issue(errors, "external-relationship", "External relationship remains in DOCX", artifact=kind)
        if not info["update_fields"]:
            add_issue(errors, "fields-not-refreshable", "DOCX does not request field refresh on open", artifact=kind)
        invalid_numpages = [value for value in info["numpages_values"] if value != str(info["pages"])]
        if invalid_numpages:
            add_issue(
                errors,
                "numpages-cache",
                "Cached NUMPAGES field does not equal the measured DOCX page count",
                artifact=kind,
                values=invalid_numpages,
                pages=info["pages"],
            )

    manual_info = artifacts["manual"][1]
    manual_cfg = config.get("manual") or {}
    structure_errors, manual_structure_stats = audit_manual_structure(manual_content, manual_cfg)
    errors.extend(structure_errors)
    non_black_text = [
        item
        for item in manual_info["text_colors"]
        if str(item.get("value") or "").upper() != "000000"
        or bool(item.get("theme"))
    ]
    if bool(manual_cfg.get("require_black_text", True)) and non_black_text:
        add_issue(
            errors,
            "manual-text-color",
            "Manual contains non-black or theme-controlled text",
            values=non_black_text[:20],
        )
    target_range = manual_cfg.get("target_page_range") or [8, 40]
    if manual_info["media_count"] < int(manual_cfg.get("minimum_images", 4)):
        add_issue(errors, "manual-images", "Manual has too few embedded images", count=manual_info["media_count"])
    minimum_runtime = int(manual_cfg.get("minimum_runtime_screenshots", 2))
    if figure_stats.get("runtime_screenshots", 0) < minimum_runtime:
        add_issue(
            errors,
            "manual-runtime-screenshots",
            "Manual has too few reviewed runtime screenshots",
            count=figure_stats.get("runtime_screenshots", 0),
            minimum=minimum_runtime,
        )
    if not (int(target_range[0]) <= manual_info["pages"] <= int(target_range[1])):
        add_issue(warnings, "manual-page-range", "Manual page count is outside the configured recommended range", pages=manual_info["pages"], target=target_range)

    application_doc = Document(args.application)
    application_info = artifacts["application"][1]
    application_cfg = config.get("application") or {}
    expected_tables = int(application_cfg.get("expected_tables", 1))
    expected_rows = int(application_cfg.get("expected_rows", 37))
    expected_columns = int(application_cfg.get("expected_columns", 13))
    structure_ok = (
        len(application_doc.tables) == expected_tables == 1
        and len(application_doc.tables[0].rows) == expected_rows
        and all(len(row.cells) == expected_columns for row in application_doc.tables[0].rows)
    )
    if not structure_ok:
        add_issue(
            errors,
            "application-structure",
            "Application draft does not preserve the configured official/reference template structure",
            expected={"tables": expected_tables, "rows": expected_rows, "columns": expected_columns},
        )
    if application_info["pages"] != int((config.get("application") or {}).get("expected_pages", 3)):
        add_issue(errors, "application-pages", "Application draft page count is not the expected value", pages=application_info["pages"])
    feature_chars = 0
    if structure_ok:
        table = application_doc.tables[0]
        if re.sub(r"\s+", "", table.cell(1, 4).text) != re.sub(r"\s+", "", name):
            add_issue(errors, "application-field", "Application key field differs from the authoritative configuration", field="software-name")
        if norm_version(table.cell(2, 12).text) != norm_version(version):
            add_issue(errors, "application-field", "Application key field differs from the authoritative configuration", field="software-version")
        if norm_date(table.cell(5, 4).text) != norm_date(completion_date):
            add_issue(errors, "application-field", "Application key field differs from the authoritative configuration", field="completion-date")
        line_cell = table.cell(23, 1).text
        found_lines = re.search(r"(?:程序量|代码)\D{0,12}(\d+)", line_cell)
        if expected_program_lines and (not found_lines or int(found_lines.group(1)) != expected_program_lines):
            add_issue(errors, "program-lines-mismatch", "Application program count differs from the frozen source line count", expected=expected_program_lines)
        feature_text = table.cell(25, 1).text
        feature_chars = len(re.sub(r"\s+", "", feature_text))
        if not 500 <= feature_chars <= 1000:
            add_issue(errors, "feature-length", "Function and technical feature text must contain 500-1000 non-whitespace characters", count=feature_chars)
        configured_feature = str(software.get("function_and_technical_features") or config.get("function_and_technical_features") or "")
        if configured_feature and re.sub(r"\s+", "", configured_feature) != re.sub(r"\s+", "", feature_text):
            add_issue(errors, "feature-text-mismatch", "Application feature text differs from the authoritative configuration")

    code_doc = Document(args.code)
    code_info = artifacts["code"][1]
    code_lines = [paragraph.text for paragraph in code_doc.paragraphs if paragraph.text.strip()]
    deposit_cfg = config.get("deposit") or {}
    long_line_threshold = int(deposit_cfg.get("long_line_warning", 120))
    if long_line_threshold <= 0:
        add_issue(errors, "deposit-long-line-threshold", "deposit.long_line_warning must be positive")
        long_line_threshold = 120
    long_line_lengths = [len(line.expandtabs(4)) for line in code_lines if len(line.expandtabs(4)) > long_line_threshold]
    if long_line_lengths:
        add_issue(
            warnings,
            "deposit-long-lines",
            "Long code lines may soft-wrap and make a 50-paragraph page display more than 50 visual lines; inspect the Office PDF",
            count=len(long_line_lengths),
            maximum=max(long_line_lengths),
            threshold=long_line_threshold,
        )
    provenance_line_count = deposit_manifest.get("deposit_line_count", deposit_manifest.get("deposit_lines"))
    if not isinstance(provenance_line_count, int) or provenance_line_count <= 0:
        add_issue(errors, "deposit-provenance-count", "Deposit provenance lacks a positive deposit_line_count")
        provenance_line_count = 0
    expected_deposit_lines = min(expected_program_lines, 3000) if expected_program_lines else provenance_line_count
    lines_per_page = int(deposit_manifest.get("lines_per_page") or deposit_cfg.get("lines_per_page", 50))
    if lines_per_page <= 0:
        add_issue(errors, "deposit-lines-per-page", "Deposit lines_per_page must be positive")
        lines_per_page = 50
    if lines_per_page != 50:
        add_issue(errors, "deposit-rule-lines", "General-deposit material must use 50 source lines per page", actual=lines_per_page)
    expected_pages = max(1, math.ceil(expected_deposit_lines / lines_per_page))
    if len(code_lines) != expected_deposit_lines:
        add_issue(errors, "deposit-line-count", "Code deposit line count does not match the source snapshot", expected=expected_deposit_lines, actual=len(code_lines))
    if code_info["pages"] != expected_pages:
        add_issue(errors, "deposit-pages", "Code DOCX page cache does not match expected pages", expected=expected_pages, actual=code_info["pages"])
    expected_breaks = max(0, expected_pages - 1)
    if code_info["page_breaks"] != expected_breaks:
        add_issue(errors, "deposit-breaks", "Code deposit page-break count is wrong", expected=expected_breaks, actual=code_info["page_breaks"])
    if not code_info["numpages_values"]:
        add_issue(errors, "deposit-numpages-field", "Code deposit has no dynamic NUMPAGES field")

    selected = deposit_manifest.get("selected")
    if not isinstance(selected, list) or len(selected) != expected_deposit_lines:
        add_issue(errors, "deposit-provenance-lines", "Deposit provenance must contain one selected record for every rendered line")
        selected = []
    if selected:
        expected_indices = (
            list(range(1, expected_program_lines + 1))
            if expected_program_lines <= 3000
            else list(range(1, 1501)) + list(range(expected_program_lines - 1499, expected_program_lines + 1))
        )
        actual_indices: list[int] = []
        for index, (line, record) in enumerate(zip(code_lines, selected), start=1):
            if not isinstance(record, dict):
                add_issue(errors, "deposit-provenance-record", "Malformed selected-line provenance record", index=index)
                continue
            try:
                actual_indices.append(int(record.get("canonical_index")))
            except (TypeError, ValueError):
                add_issue(errors, "deposit-provenance-index", "Selected-line canonical index is missing", index=index)
            if as_int(record.get("deposit_index")) != index:
                add_issue(errors, "deposit-provenance-order", "Selected-line deposit_index is not continuous", index=index)
            if as_int(record.get("page")) != ((index - 1) // lines_per_page) + 1 or as_int(record.get("line_on_page")) != ((index - 1) % lines_per_page) + 1:
                add_issue(errors, "deposit-provenance-page", "Selected-line page provenance is inconsistent", index=index)
            digest = hashlib.sha256(line.encode("utf-8")).hexdigest()
            if digest != record.get("rendered_line_sha256"):
                add_issue(errors, "deposit-line-content", "Rendered code line differs from deposit provenance", index=index)
        if actual_indices != expected_indices:
            add_issue(errors, "deposit-selection", "Deposit is not the full source or reproducible first/last 1500-line selection")
        text_digest = hashlib.sha256()
        for line in code_lines:
            text_digest.update(line.encode("utf-8"))
            text_digest.update(b"\n")
        if text_digest.hexdigest() != deposit_manifest.get("deposit_text_sha256"):
            add_issue(errors, "deposit-text-hash", "Code deposit text SHA-256 differs from provenance")

    provenance_sources = deposit_manifest.get("source_files")
    if not isinstance(provenance_sources, list):
        add_issue(errors, "deposit-source-provenance", "Deposit provenance lacks source_files")
    elif manifest:
        frozen = [item for item in manifest.get("files", []) if isinstance(item, dict)]
        if len(provenance_sources) != len(frozen):
            add_issue(errors, "deposit-source-count", "Deposit provenance and frozen manifest contain different source file counts")
        else:
            for expected, actual in zip(frozen, provenance_sources):
                if not isinstance(actual, dict) or (
                    actual.get("path") != expected.get("path")
                    or str(actual.get("sha256", "")).lower() != str(expected.get("sha256", "")).lower()
                    or as_int(actual.get("effective_lines")) != as_int(expected.get("nonblank_lines"), -2)
                ):
                    add_issue(errors, "deposit-source-mismatch", "Deposit provenance differs from the frozen source manifest", file=expected.get("path"))
                    break
    repeated_10 = [positions for positions in normalized_windows(code_lines, 10).values() if len(positions) > 1]
    if repeated_10:
        add_issue(errors, "deposit-repeat-10", "Exact normalized 10-line blocks repeat inside the code deposit", groups=len(repeated_10), positions=repeated_10[:20])
    repeated_5 = [positions for positions in normalized_windows(code_lines, 5).values() if len(positions) > 1]
    if repeated_5:
        add_issue(warnings, "deposit-repeat-5", "Five-line repeated blocks require human classification", groups=len(repeated_5))
    if deposit_manifest.get("snapshot_sha256") and manifest.get("snapshot_sha256") and deposit_manifest["snapshot_sha256"] != manifest["snapshot_sha256"]:
        add_issue(errors, "deposit-snapshot", "Deposit was built from a different source snapshot")
    provenance_software = deposit_manifest.get("software") or {}
    if provenance_software and (
        str(provenance_software.get("name") or "").strip() != name
        or norm_version(provenance_software.get("version")) != norm_version(version)
    ):
        add_issue(errors, "deposit-software-identity", "Deposit provenance uses a different software name or version")

    pdf_counts: dict[str, int | None] = {}
    if args.pdf_dir:
        for kind, (path, info) in artifacts.items():
            pdf = args.pdf_dir / f"{path.stem}.pdf"
            count = pdf_page_count(pdf)
            pdf_counts[kind] = count
            if count is None:
                add_issue(errors, "pdf-missing", "Rendered PDF preview is missing or unreadable", artifact=kind)
            elif count != info["pages"]:
                add_issue(errors, "pdf-page-mismatch", "DOCX cached pages and PDF pages differ", artifact=kind, docx=info["pages"], pdf=count)
    else:
        add_issue(errors, "pdf-not-checked", "No Office-rendered PDF preview directory was supplied; pagination is unverified")

    unique_versions = {norm_version(version)} if version else set()
    unique_versions.update(norm_version(value) for value in facts.get("version_values", []) if value)
    if len(unique_versions) > 1:
        add_issue(errors, "cross-material-version", "Project and registered version values are inconsistent", values=sorted(unique_versions))

    report = {
        "schema_version": 1,
        "submission_ready": not errors,
        "errors": errors,
        "warnings": warnings,
        "software": {"full_name": name, "version": version, "completion_date": completion_date},
        "artifacts": {
            kind: {
                "path": str(path),
                "pages": info["pages"],
                "media_count": info["media_count"],
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            }
            for kind, (path, info) in artifacts.items()
        },
        "application_feature_chars": feature_chars,
        "code_nonblank_paragraphs": len(code_lines),
        "source_snapshot": snapshot_stats,
        "feature_evidence": evidence_stats,
        "figures": figure_stats,
        "manual_structure": manual_structure_stats,
        "pdf_pages": pdf_counts,
        "official_form_note": "This Word application is a drafting aid. Use the current receiving platform's official form when required.",
        "external_similarity_note": "Local checks cannot guarantee non-overlap with an unknown registration-agency code corpus.",
    }
    args.report_json.parent.mkdir(parents=True, exist_ok=True)
    args.report_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    if args.report_md:
        status = "可进入人工提交复核" if report["submission_ready"] else "不可提交"
        lines = [f"# 软著三件套终检：{status}", "", f"- 软件：{name}", f"- 版本：{version}", f"- 错误：{len(errors)}", f"- 警告：{len(warnings)}", "", "## 阻断项", ""]
        lines.extend(f"- `{item['code']}`：{item['message']}" for item in errors)
        lines.extend(["", "## 警告", ""])
        lines.extend(f"- `{item['code']}`：{item['message']}" for item in warnings)
        args.report_md.parent.mkdir(parents=True, exist_ok=True)
        args.report_md.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(json.dumps({"submission_ready": report["submission_ready"], "errors": len(errors), "warnings": len(warnings), "report": str(args.report_json)}, ensure_ascii=False))
    return 0 if report["submission_ready"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
