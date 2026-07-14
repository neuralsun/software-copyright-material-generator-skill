#!/usr/bin/env python3
"""Build a traceable Chinese software-copyright source-code deposit DOCX.

The input is a JSON manifest.  Explicit ``source.files`` order is treated as
the canonical source-program order.  When it is omitted, deterministic glob
discovery is used.  Source text is decoded strictly; this tool never replaces
undecodable bytes silently.
"""

from __future__ import annotations

import argparse
import fnmatch
import hashlib
import json
import math
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SCHEMA_VERSION = "1.0"
DEFAULT_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030")
DEFAULT_EXTENSIONS = {
    ".c", ".cc", ".cpp", ".cs", ".css", ".go", ".h", ".hpp", ".html",
    ".java", ".js", ".jsx", ".kt", ".less", ".m", ".mm", ".php", ".py",
    ".rb", ".rs", ".scss", ".sql", ".svelte", ".swift", ".ts", ".tsx",
    ".vue", ".xml",
}
DEFAULT_EXCLUDED_DIRS = {
    ".agents", ".cache", ".git", ".hg", ".idea", ".mypy_cache", ".next",
    ".pytest_cache", ".ruff_cache", ".runtime", ".runtime-logs", ".svn",
    ".venv", ".vscode", "__pycache__", "bin", "build", "copyright-work",
    "coverage", "dist", "docs_gpt", "materials", "media", "node_modules",
    "obj", "previews", "screenshots", "target", "temp", "tmp", "uploads",
    "vendor", "venv", "材料",
}
DEFAULT_EXCLUDE_GLOBS = (
    "**/*.min.js", "**/*.min.css", "**/*.map", "**/*.lock", "**/.env",
    "**/.env.*", "**/package-lock.json", "**/pnpm-lock.yaml", "**/yarn.lock",
)


class DepositError(RuntimeError):
    """Raised for a manifest, source, or output validation failure."""


@dataclass(frozen=True)
class SourceLine:
    canonical_index: int
    path: str
    source_line: int
    text: str
    encoding: str

    def location(self) -> dict[str, Any]:
        return {
            "canonical_index": self.canonical_index,
            "path": self.path,
            "source_line": self.source_line,
        }


@dataclass(frozen=True)
class SourceFileInfo:
    path: str
    encoding: str
    bytes: int
    physical_lines: int
    effective_lines: int
    sha256: str
    role: str | None = None


def load_json_manifest(argument: str) -> tuple[dict[str, Any], Path]:
    """Load a UTF-8 JSON manifest and return it with its path-resolution base."""
    try:
        if argument == "-":
            data = json.load(sys.stdin)
            base = Path.cwd()
        else:
            path = Path(argument).expanduser().resolve()
            data = json.loads(path.read_text(encoding="utf-8-sig"))
            base = path.parent
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise DepositError(f"无法读取JSON manifest：{exc}") from exc
    if not isinstance(data, dict):
        raise DepositError("manifest顶层必须是JSON对象")
    source_manifest_value = data.get("source_manifest")
    if source_manifest_value is not None:
        if not isinstance(source_manifest_value, str) or not source_manifest_value.strip():
            raise DepositError("source_manifest必须是非空JSON路径")
        source_manifest_path = resolve_from_base(source_manifest_value, base)
        try:
            frozen = json.loads(source_manifest_path.read_text(encoding="utf-8-sig"))
        except (OSError, UnicodeError, json.JSONDecodeError) as exc:
            raise DepositError(f"无法读取source_manifest：{exc}") from exc
        if not isinstance(frozen, dict):
            raise DepositError("source_manifest顶层必须是JSON对象")
        # Wrapper values (software/deposit/audit/output) win; the frozen source
        # contributes project_root, ordered files, counts, and snapshot hashes.
        merged = dict(frozen)
        merged.update(data)
        wrapper_source = data.get("source", data.get("sources"))
        wrapper_has_files = "files" in data or (
            isinstance(wrapper_source, dict) and "files" in wrapper_source
        )
        if not wrapper_has_files:
            merged["files"] = frozen.get("files", [])
        data = merged
    return data, base


def _as_dict(value: Any, label: str) -> dict[str, Any]:
    if value is None:
        return {}
    if not isinstance(value, dict):
        raise DepositError(f"{label}必须是JSON对象")
    return value


def _as_string_list(value: Any, label: str) -> list[str]:
    if value is None:
        return []
    if not isinstance(value, list) or not all(isinstance(item, str) for item in value):
        raise DepositError(f"{label}必须是字符串数组")
    return value


def resolve_from_base(value: str | Path, base: Path) -> Path:
    path = Path(value).expanduser()
    return (path if path.is_absolute() else base / path).resolve()


def project_root_from_manifest(manifest: dict[str, Any], base: Path) -> Path:
    value = manifest.get("project_root")
    if value is None:
        value = _as_dict(manifest.get("project"), "project").get("root", ".")
    if not isinstance(value, str) or not value.strip():
        raise DepositError("project_root必须是非空字符串")
    root = resolve_from_base(value, base)
    if not root.is_dir():
        raise DepositError(f"项目根目录不存在：{root}")
    return root


def source_config_from_manifest(manifest: dict[str, Any]) -> dict[str, Any]:
    value = manifest.get("source")
    if value is None:
        value = manifest.get("sources")
    return _as_dict(value, "source")


def _relative_source_path(path: Path, root: Path) -> str:
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError as exc:
        raise DepositError(f"源码文件越出project_root：{path}") from exc


def _matches_any(path: str, patterns: Iterable[str]) -> bool:
    normalized = path.replace("\\", "/")
    return any(
        fnmatch.fnmatchcase(normalized, pattern.replace("\\", "/"))
        for pattern in patterns
    )


def _path_is_excluded(relative: str, exclude_globs: Sequence[str]) -> bool:
    parts = {part.lower() for part in Path(relative).parts}
    return bool(parts & DEFAULT_EXCLUDED_DIRS) or _matches_any(relative, exclude_globs)


def collect_source_entries(
    manifest: dict[str, Any], root: Path
) -> list[dict[str, Any]]:
    """Return ordered source entries from explicit files or deterministic discovery."""
    config = source_config_from_manifest(manifest)
    raw_files = config.get("files")
    if raw_files is None:
        raw_files = manifest.get("source_files")
    if raw_files is None and "files" in manifest:
        raw_files = manifest.get("files")

    exclude_globs = list(DEFAULT_EXCLUDE_GLOBS)
    exclude_globs.extend(_as_string_list(config.get("exclude_globs"), "source.exclude_globs"))
    entries: list[dict[str, Any]] = []

    if raw_files is not None:
        if not isinstance(raw_files, list):
            raise DepositError("source.files必须是数组")
        for position, item in enumerate(raw_files, start=1):
            if isinstance(item, str):
                entry: dict[str, Any] = {"path": item}
            elif isinstance(item, dict):
                entry = dict(item)
            else:
                raise DepositError(f"source.files[{position - 1}]必须是字符串或对象")
            if entry.get("include", True) is False:
                continue
            value = entry.get("path")
            if not isinstance(value, str) or not value.strip():
                raise DepositError(f"source.files[{position - 1}].path缺失")
            path = resolve_from_base(value, root)
            relative = _relative_source_path(path, root)
            if _path_is_excluded(relative, exclude_globs) and not entry.get("force_include", False):
                raise DepositError(f"显式源码命中排除规则：{relative}；如确需纳入请设置force_include=true")
            entry["absolute_path"] = path
            entry["relative_path"] = relative
            entries.append(entry)
    else:
        include_globs = _as_string_list(config.get("include_globs"), "source.include_globs")
        extensions = {
            str(item).lower() if str(item).startswith(".") else f".{str(item).lower()}"
            for item in config.get("extensions", sorted(DEFAULT_EXTENSIONS))
        }
        discovered: dict[str, Path] = {}
        if include_globs:
            for pattern in include_globs:
                if Path(pattern).is_absolute() or ".." in Path(pattern).parts:
                    raise DepositError(f"include_globs不得越出项目根目录：{pattern}")
                for path in root.glob(pattern):
                    if path.is_file():
                        relative = _relative_source_path(path, root)
                        if not _path_is_excluded(relative, exclude_globs):
                            discovered[relative] = path.resolve()
        else:
            for path in root.rglob("*"):
                if not path.is_file() or path.suffix.lower() not in extensions:
                    continue
                relative = _relative_source_path(path, root)
                if not _path_is_excluded(relative, exclude_globs):
                    discovered[relative] = path.resolve()
        entries = [
            {"path": relative, "absolute_path": discovered[relative], "relative_path": relative}
            for relative in sorted(discovered, key=lambda item: (item.lower(), item))
        ]

    if not entries:
        raise DepositError("未找到可交存的第一方源码文件")
    duplicates = [path for path, count in Counter(e["relative_path"] for e in entries).items() if count > 1]
    if duplicates:
        raise DepositError(f"source.files存在重复文件：{duplicates}")
    return entries


def decode_source_bytes(
    raw: bytes,
    *,
    path: Path,
    declared_encoding: str | None,
    candidates: Sequence[str],
) -> tuple[str, str]:
    if raw.startswith((b"\xff\xfe", b"\xfe\xff")):
        ordered = [declared_encoding] if declared_encoding else ["utf-16"]
    elif raw.startswith(b"\xef\xbb\xbf"):
        ordered = [declared_encoding] if declared_encoding else ["utf-8-sig"]
    else:
        ordered = [declared_encoding] if declared_encoding else list(candidates)
    ordered = [item for index, item in enumerate(ordered) if item and item not in ordered[:index]]
    errors: list[str] = []
    for encoding in ordered:
        try:
            text = raw.decode(encoding, errors="strict")
        except (LookupError, UnicodeDecodeError) as exc:
            errors.append(f"{encoding}: {exc}")
            continue
        if "\ufffd" in text:
            errors.append(f"{encoding}: 包含Unicode替换字符U+FFFD")
            continue
        if "\x00" in text:
            errors.append(f"{encoding}: 解码结果包含NUL，疑似二进制文件")
            continue
        return text, encoding
    detail = "; ".join(errors[:4])
    raise DepositError(f"源码无法严格解码：{path}（{detail}）")


def load_project_source(
    manifest: dict[str, Any], base: Path
) -> tuple[list[SourceLine], list[SourceFileInfo], Path]:
    root = project_root_from_manifest(manifest, base)
    config = source_config_from_manifest(manifest)
    candidates = tuple(_as_string_list(config.get("encoding_candidates"), "source.encoding_candidates"))
    if not candidates:
        candidates = DEFAULT_ENCODINGS
    tab_size = int(config.get("tab_size", 4))
    max_file_bytes = int(config.get("max_file_bytes", 20 * 1024 * 1024))
    if tab_size < 1 or tab_size > 16:
        raise DepositError("source.tab_size必须在1—16之间")
    if max_file_bytes < 1:
        raise DepositError("source.max_file_bytes必须为正数")

    lines: list[SourceLine] = []
    files: list[SourceFileInfo] = []
    for entry in collect_source_entries(manifest, root):
        path: Path = entry["absolute_path"]
        relative: str = entry["relative_path"]
        if not path.is_file():
            raise DepositError(f"源码文件不存在：{path}")
        raw = path.read_bytes()
        if len(raw) > max_file_bytes:
            raise DepositError(f"单个源码文件超过大小上限：{relative}（{len(raw)} bytes）")
        declared = entry.get("encoding")
        if declared is not None and not isinstance(declared, str):
            raise DepositError(f"encoding必须是字符串：{relative}")
        text, encoding = decode_source_bytes(
            raw, path=path, declared_encoding=declared, candidates=candidates
        )
        actual_digest = hashlib.sha256(raw).hexdigest()
        expected_digest = entry.get("sha256")
        if expected_digest and str(expected_digest).lower() != actual_digest:
            raise DepositError(f"源码在manifest冻结后发生变化：{relative}（SHA-256不一致）")
        physical = text.splitlines()
        effective_count = 0
        for number, raw_line in enumerate(physical, start=1):
            rendered = raw_line.expandtabs(tab_size).rstrip()
            if not rendered.strip():
                continue
            effective_count += 1
            lines.append(
                SourceLine(
                    canonical_index=len(lines) + 1,
                    path=relative,
                    source_line=number,
                    text=rendered,
                    encoding=encoding,
                )
            )
        expected_physical = entry.get("physical_lines")
        if expected_physical is not None and int(expected_physical) != len(physical):
            raise DepositError(f"源码在manifest冻结后发生变化：{relative}（物理行数不一致）")
        expected_effective = entry.get("nonblank_lines", entry.get("effective_lines"))
        if expected_effective is not None and int(expected_effective) != effective_count:
            raise DepositError(f"源码在manifest冻结后发生变化：{relative}（非空行数不一致）")
        files.append(
            SourceFileInfo(
                path=relative,
                encoding=encoding,
                bytes=len(raw),
                physical_lines=len(physical),
                effective_lines=effective_count,
                sha256=actual_digest,
                role=str(entry.get("role")) if entry.get("role") is not None else None,
            )
        )
    if not lines:
        raise DepositError("源码清单没有任何非空源码行")
    expected_program_lines = manifest.get("program_nonblank_lines")
    if expected_program_lines is not None and int(expected_program_lines) != len(lines):
        raise DepositError("完整源码非空行数与冻结manifest不一致")
    expected_snapshot = manifest.get("snapshot_sha256")
    if expected_snapshot:
        scanner_digest = hashlib.sha256()
        for info in files:
            scanner_digest.update(
                f"{info.path}\0{info.sha256}\0{info.effective_lines}\n".encode("utf-8")
            )
        if scanner_digest.hexdigest().lower() != str(expected_snapshot).lower():
            raise DepositError("完整源码快照与冻结manifest的SHA-256不一致")
    return lines, files, root


def select_deposit_lines(
    lines: Sequence[SourceLine], *, lines_per_page: int, front_pages: int, back_pages: int
) -> tuple[list[SourceLine], str]:
    if min(lines_per_page, front_pages, back_pages) <= 0:
        raise DepositError("lines_per_page、front_pages、back_pages必须为正整数")
    front_count = lines_per_page * front_pages
    back_count = lines_per_page * back_pages
    threshold = front_count + back_count
    if len(lines) <= threshold:
        mode = "full_source" if len(lines) < threshold else "full_source_at_threshold"
        return list(lines), mode
    return list(lines[:front_count]) + list(lines[-back_count:]), "front_and_back"


def normalized_line(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _trivial_line(text: str) -> bool:
    value = normalized_line(text)
    if not value:
        return True
    if re.fullmatch(r"[{}\[\]();,.:<>/\\]+", value):
        return True
    if re.fullmatch(r"</?[A-Za-z][A-Za-z0-9:_-]*\s*/?>", value):
        return True
    return False


def analyze_repeated_blocks(
    lines: Sequence[SourceLine], widths: Sequence[int], *, max_examples: int = 20
) -> list[dict[str, Any]]:
    normalized = [normalized_line(item.text) for item in lines]
    results: list[dict[str, Any]] = []
    for width in sorted(set(int(value) for value in widths if int(value) > 1)):
        occurrences: dict[str, list[int]] = defaultdict(list)
        for index in range(max(0, len(normalized) - width + 1)):
            key = "\n".join(normalized[index : index + width])
            occurrences[key].append(index)
        duplicates = [(key, positions) for key, positions in occurrences.items() if len(positions) > 1]
        duplicates.sort(key=lambda item: (-len(item[1]), item[1][0]))
        examples: list[dict[str, Any]] = []
        for key, positions in duplicates[:max_examples]:
            examples.append(
                {
                    "occurrences": [lines[index].location() for index in positions[:10]],
                    "occurrence_count": len(positions),
                    "trivial": all(_trivial_line(value) for value in key.splitlines()),
                    "preview": key.splitlines()[: min(width, 12)],
                }
            )
        results.append(
            {
                "width": width,
                "unique_duplicate_blocks": len(duplicates),
                "duplicate_window_count": sum(len(positions) - 1 for _, positions in duplicates),
                "examples": examples,
                "examples_truncated": len(duplicates) > max_examples,
            }
        )
    return results


SENSITIVE_RULES: tuple[tuple[str, str, str], ...] = (
    ("private_key", "critical", r"-----BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY-----"),
    ("aws_access_key", "critical", r"\bAKIA[0-9A-Z]{16}\b"),
    ("jwt_token", "high", r"\beyJ[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{8,}\b"),
    (
        "mojibake",
        "high",
        r"[�\ue000-\uf8ff]|(?:[鏈鐢璁鍖缃锛銆搴櫌涓鍦鐨鏄埛綍嶅瓨ㄣ].{0,3}){3,}",
    ),
    (
        "credential_assignment",
        "high",
        r"(?i)(?:api[_-]?key|access[_-]?key|access[_-]?token|auth[_-]?token|"
        r"client[_-]?secret|password|passwd|pwd|secret(?:[_-]?key)?|token|"
        r"invite[_-]?code|session[_-]?key)[\"']?\s*"
        r"(?:(?::\s*[A-Za-z_][\w.\[\]|, ]*\s*=)|[:=])\s*[rubfRUBF]*"
        r"([\"'])(?!\s*(?:none|null|your[_-]|example|placeholder))[^\"'\r\n]{6,}\1",
    ),
    (
        "environment_secret_default",
        "high",
        r"(?i)(?:os\.)?(?:getenv|environ\.get)\(\s*[\"'][^\"']*"
        r"(?:key|secret|password|passwd|pwd|token|invite)[^\"']*[\"']\s*,\s*"
        r"[\"'](?!\s*(?:none|null|your[_-]|example|placeholder))[^\"'\r\n]{6,}[\"']",
    ),
    (
        "password_like_literal",
        "high",
        r"[\"'][A-Za-z][A-Za-z0-9_@.\-]*(?:12345678|123456|654321)[\"']",
    ),
    (
        "credential_uri",
        "high",
        r"(?i)\b(?:postgres(?:ql)?|mysql|mongodb(?:\+srv)?|redis)://[^\s:/@]+:[^\s/@]+@",
    ),
    ("cn_identity_candidate", "medium", r"(?<!\d)[1-9]\d{5}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[0-9Xx](?!\d)"),
    ("cn_mobile_candidate", "medium", r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    ("email_candidate", "low", r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b"),
)


def _redacted_preview(text: str, limit: int = 220) -> str:
    value = text.strip()[:limit]
    value = re.sub(r"([\"'])[^\"']{5,}\1", r"\1<redacted>\1", value)
    value = re.sub(r"(?i)(://[^\s:/@]+:)[^\s/@]+(@)", r"\1<redacted>\2", value)
    value = re.sub(r"\beyJ[A-Za-z0-9_.-]{20,}\b", "<redacted-jwt>", value)
    return value


def scan_sensitive_lines(lines: Sequence[SourceLine], *, max_findings: int = 200) -> list[dict[str, Any]]:
    compiled = [(name, severity, re.compile(pattern)) for name, severity, pattern in SENSITIVE_RULES]
    findings: list[dict[str, Any]] = []
    for item in lines:
        matched_rules: set[str] = set()
        for name, severity, pattern in compiled:
            if pattern.search(item.text):
                if name == "password_like_literal" and matched_rules & {
                    "credential_assignment",
                    "environment_secret_default",
                }:
                    continue
                findings.append(
                    {
                        "rule": name,
                        "severity": severity,
                        **item.location(),
                        "preview": _redacted_preview(item.text),
                    }
                )
                matched_rules.add(name)
                if len(findings) >= max_findings:
                    return findings
    return findings


LOGIN_PATH_RE = re.compile(r"(?i)(?:^|[/_.-])(?:login|signin|sign-in|register|signup|auth)(?:[/_.-]|$)")
LOGIN_INDICATORS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("password_input", re.compile(r"(?i)(?:type\s*=\s*[\"']password|password|passwd)")),
    ("form_markup", re.compile(r"(?i)<form\b")),
    ("html_shell", re.compile(r"(?i)(?:<!doctype\s+html|<html\b|<head\b)")),
    ("login_word", re.compile(r"(?i)(?:\blog[ -]?in\b|\bsign[ -]?in\b|登录|登陆)")),
    ("register_word", re.compile(r"(?i)(?:\bregister\b|\bsign[ -]?up\b|注册)")),
    ("remember_me", re.compile(r"(?i)(?:remember\s+me|记住我)")),
)


def scan_login_risk(lines: Sequence[SourceLine]) -> list[dict[str, Any]]:
    grouped: dict[str, list[SourceLine]] = defaultdict(list)
    for item in lines:
        grouped[item.path].append(item)
    findings: list[dict[str, Any]] = []
    ui_extensions = {".css", ".html", ".htm", ".js", ".jsx", ".scss", ".svelte", ".ts", ".tsx", ".vue"}
    for path, items in grouped.items():
        joined = "\n".join(item.text for item in items)
        indicators = [name for name, pattern in LOGIN_INDICATORS if pattern.search(joined)]
        path_match = bool(LOGIN_PATH_RE.search(path))
        is_ui = Path(path).suffix.lower() in ui_extensions
        if not path_match and len(indicators) < 2:
            continue
        if is_ui and "password_input" in indicators and ("form_markup" in indicators or "login_word" in indicators):
            severity = "high"
            reason = "交存片段包含通用登录表单表现层，可能与常见模板相似"
        elif is_ui and (path_match or "html_shell" in indicators):
            severity = "medium"
            reason = "交存片段包含登录/注册或入口页面表现层"
        else:
            severity = "info"
            reason = "交存片段包含认证业务实现；该类服务端逻辑通常是功能证据"
        findings.append(
            {
                "path": path,
                "severity": severity,
                "reason": reason,
                "path_pattern_match": path_match,
                "indicators": indicators,
                "selected_line_count": len(items),
                "first_source_line": items[0].source_line,
                "last_source_line": items[-1].source_line,
            }
        )
    return sorted(findings, key=lambda item: ({"high": 0, "medium": 1, "info": 2}[item["severity"]], item["path"]))


def selected_forbidden_paths(lines: Sequence[SourceLine], patterns: Sequence[str]) -> list[str]:
    paths = sorted({item.path for item in lines})
    return [path for path in paths if _matches_any(path, patterns)]


def snapshot_hash(lines: Sequence[SourceLine], *, include_provenance: bool) -> str:
    digest = hashlib.sha256()
    for item in lines:
        if include_provenance:
            digest.update(item.path.encode("utf-8"))
            digest.update(b"\0")
            digest.update(str(item.source_line).encode("ascii"))
            digest.update(b"\0")
        digest.update(item.text.encode("utf-8"))
        digest.update(b"\n")
    return digest.hexdigest()


def set_run_font(run: Any, *, ascii_font: str, east_asia_font: str, size: float, bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)


def add_complex_field(
    paragraph: Any,
    instruction: str,
    fallback: str,
    *,
    ascii_font: str,
    east_asia_font: str,
    size: float,
    bold: bool = False,
) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction.strip()} "
    instruction_run._r.append(instr)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    result_run = paragraph.add_run(fallback)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    for run in (begin_run, instruction_run, separate_run, result_run, end_run):
        set_run_font(
            run,
            ascii_font=ascii_font,
            east_asia_font=east_asia_font,
            size=size,
            bold=bold,
        )


def enable_field_updates(document: Document) -> None:
    settings = document.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def set_continuous_line_numbering(section: Any) -> None:
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:lnNumType"))
    if old is not None:
        sect_pr.remove(old)
    node = OxmlElement("w:lnNumType")
    node.set(qn("w:countBy"), "1")
    node.set(qn("w:restart"), "continuous")
    node.set(qn("w:distance"), "240")
    sect_pr.append(node)


def _float_config(config: dict[str, Any], name: str, default: float) -> float:
    try:
        value = float(config.get(name, default))
    except (TypeError, ValueError) as exc:
        raise DepositError(f"deposit.{name}必须是数字") from exc
    if value <= 0:
        raise DepositError(f"deposit.{name}必须为正数")
    return value


def build_code_docx(
    selected: Sequence[SourceLine],
    output: Path,
    *,
    software_name: str,
    version: str,
    deposit_config: dict[str, Any],
    expected_pages: int,
) -> None:
    fonts = _as_dict(deposit_config.get("fonts"), "deposit.fonts")
    ascii_font = str(fonts.get("ascii", deposit_config.get("font_ascii", "Arial Narrow")))
    east_font = str(fonts.get("east_asia", deposit_config.get("font_east_asia", "宋体")))
    font_size = _float_config(deposit_config, "font_size_pt", 10.5)
    line_spacing = _float_config(deposit_config, "line_spacing_pt", 12.0)
    lines_per_page = int(deposit_config.get("lines_per_page", 50))

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(_float_config(deposit_config, "page_width_cm", 21.001))
    section.page_height = Cm(_float_config(deposit_config, "page_height_cm", 29.704))
    section.top_margin = Cm(_float_config(deposit_config, "top_margin_cm", 2.491))
    section.bottom_margin = Cm(_float_config(deposit_config, "bottom_margin_cm", 2.011))
    section.left_margin = Cm(_float_config(deposit_config, "left_margin_cm", 2.491))
    section.right_margin = Cm(_float_config(deposit_config, "right_margin_cm", 2.491))
    section.header_distance = Cm(_float_config(deposit_config, "header_distance_cm", 1.27))
    section.footer_distance = Cm(_float_config(deposit_config, "footer_distance_cm", 1.27))
    set_continuous_line_numbering(section)

    normal = document.styles["Normal"]
    normal.font.name = ascii_font
    normal.font.size = Pt(font_size)
    normal_fonts = normal._element.get_or_add_rPr().rFonts
    normal_fonts.set(qn("w:ascii"), ascii_font)
    normal_fonts.set(qn("w:hAnsi"), ascii_font)
    normal_fonts.set(qn("w:eastAsia"), east_font)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(line_spacing)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.clear()
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(0)
    usable_width = section.page_width - section.left_margin - section.right_margin
    header_p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    title_run = header_p.add_run(f"{software_name} {version}")
    set_run_font(title_run, ascii_font=ascii_font, east_asia_font=east_font, size=10, bold=True)
    tab_run = header_p.add_run("\t")
    set_run_font(tab_run, ascii_font=ascii_font, east_asia_font=east_font, size=9)
    add_complex_field(
        header_p,
        "PAGE",
        "1",
        ascii_font=ascii_font,
        east_asia_font=east_font,
        size=9,
        bold=True,
    )

    footer_p = section.footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    add_complex_field(
        footer_p,
        "PAGE",
        "1",
        ascii_font=ascii_font,
        east_asia_font=east_font,
        size=8,
        bold=True,
    )
    of_run = footer_p.add_run(" of ")
    set_run_font(of_run, ascii_font=ascii_font, east_asia_font=east_font, size=8)
    add_complex_field(
        footer_p,
        "NUMPAGES",
        str(expected_pages),
        ascii_font=ascii_font,
        east_asia_font=east_font,
        size=8,
        bold=True,
    )
    enable_field_updates(document)

    for index, item in enumerate(selected, start=1):
        paragraph = document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_spacing)
        fmt.keep_together = False
        fmt.keep_with_next = False
        fmt.widow_control = False
        run = paragraph.add_run(item.text)
        set_run_font(run, ascii_font=ascii_font, east_asia_font=east_font, size=font_size)
        if index % lines_per_page == 0 and index < len(selected):
            run.add_break(WD_BREAK.PAGE)

    now = datetime.now(timezone.utc).replace(tzinfo=None)
    properties = document.core_properties
    properties.title = f"{software_name} {version} 源程序一般交存文档"
    properties.subject = "源程序一般交存鉴别材料"
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.created = now
    properties.modified = now

    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.stem}.tmp.docx")
    if temporary.exists():
        temporary.unlink()
    document.save(temporary)
    temporary.replace(output)


def atomic_write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    temporary.replace(path)


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[<>:\"/\\|?*\x00-\x1f]", "_", value).strip().rstrip(".")
    return cleaned or "软件"


def _output_path(value: str | None, default: Path, base: Path) -> Path:
    return resolve_from_base(value, base) if value else default.resolve()


def _severity_at_least(value: str, threshold: str) -> bool:
    order = {"info": 0, "low": 1, "medium": 2, "high": 3, "critical": 4}
    return order.get(value, 0) >= order.get(threshold, 3)


def run_build(args: argparse.Namespace) -> int:
    manifest, base = load_json_manifest(args.manifest)
    software = _as_dict(manifest.get("software"), "software")
    software_name = software.get(
        "name", software.get("full_name", software.get("expected_name", manifest.get("software_name")))
    )
    version = software.get("version", software.get("expected_version", manifest.get("version")))
    if not isinstance(software_name, str) or not software_name.strip():
        raise DepositError("software.name/full_name不能为空")
    if not isinstance(version, str) or not version.strip():
        raise DepositError("software.version不能为空")
    software_name, version = software_name.strip(), version.strip()

    deposit = _as_dict(manifest.get("deposit"), "deposit")
    audit = _as_dict(manifest.get("audit"), "audit")
    lines_per_page = int(deposit.get("lines_per_page", 50))
    front_pages = int(deposit.get("front_pages", 30))
    back_pages = int(deposit.get("back_pages", 30))
    if min(lines_per_page, front_pages, back_pages) <= 0:
        raise DepositError("交存分页参数必须为正整数")

    full_lines, file_infos, project_root = load_project_source(manifest, base)
    selected, selection_mode = select_deposit_lines(
        full_lines,
        lines_per_page=lines_per_page,
        front_pages=front_pages,
        back_pages=back_pages,
    )
    expected_pages = math.ceil(len(selected) / lines_per_page)
    widths = audit.get("repeat_widths", [5, 9, 10, 20])
    if not isinstance(widths, list):
        raise DepositError("audit.repeat_widths必须是整数数组")
    repeats = analyze_repeated_blocks(selected, [int(value) for value in widths])
    sensitive = scan_sensitive_lines(selected)
    login_risks = scan_login_risk(selected)
    forbidden_patterns = _as_string_list(
        audit.get("forbidden_selected_paths"), "audit.forbidden_selected_paths"
    )
    forbidden = selected_forbidden_paths(selected, forbidden_patterns)

    configured_output = deposit.get("output_docx")
    default_output = base / f"{_safe_filename(software_name)}-代码(一般交存版).docx"
    output = _output_path(args.output or configured_output, default_output, base)
    provenance_output = _output_path(
        args.provenance or deposit.get("provenance_output"),
        output.with_suffix(".provenance.json"),
        base,
    )
    report_output = _output_path(
        args.report or deposit.get("report_output"),
        output.with_suffix(".audit.json"),
        base,
    )

    high_sensitive = [item for item in sensitive if _severity_at_least(item["severity"], "high")]
    high_login = [item for item in login_risks if item["severity"] == "high"]
    blockers: list[str] = []
    if forbidden:
        blockers.append(f"交存摘录命中禁止路径：{forbidden}")
    if high_sensitive and not args.allow_sensitive and audit.get("block_on_sensitive", True):
        blockers.append(f"发现{len(high_sensitive)}项高风险敏感信息候选")
    if high_login and (args.fail_on_login_risk or audit.get("block_on_high_login_risk", False)):
        blockers.append(f"发现{len(high_login)}个高风险登录表现层文件")

    long_line_threshold = int(audit.get("long_line_warning", 120))
    line_lengths = [len(item.text) for item in selected]
    warnings: list[str] = []
    if any(length > long_line_threshold for length in line_lengths):
        warnings.append(
            f"{sum(length > long_line_threshold for length in line_lengths)}条代码超过"
            f"{long_line_threshold}字符，必须用Word/WPS核验换行后实际页数"
        )
    if len(full_lines) < lines_per_page * (front_pages + back_pages):
        warnings.append("完整源码不足60页，已按一般交存规则纳入全部源码，未复制凑页")
    if lines_per_page != 50 or front_pages != 30 or back_pages != 30:
        warnings.append("当前分页参数不是常用的一般交存50条/页、前后各30页，请核对适用规则")

    selected_entries = []
    for index, item in enumerate(selected, start=1):
        selected_entries.append(
            {
                "deposit_index": index,
                "page": ((index - 1) // lines_per_page) + 1,
                "line_on_page": ((index - 1) % lines_per_page) + 1,
                **item.location(),
                "rendered_line_sha256": hashlib.sha256(item.text.encode("utf-8")).hexdigest(),
            }
        )
    provenance = {
        "schema_version": SCHEMA_VERSION,
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "project_root": str(project_root),
        "software": {"name": software_name, "version": version},
        "counting_method": "非空物理行；Tab按配置展开；仅去除行尾空白；严格解码",
        "source_files": [info.__dict__ for info in file_infos],
        "full_source_line_count": len(full_lines),
        "full_snapshot_sha256": snapshot_hash(full_lines, include_provenance=True),
        "selection_mode": selection_mode,
        "deposit_line_count": len(selected),
        "deposit_text_sha256": snapshot_hash(selected, include_provenance=False),
        "lines_per_page": lines_per_page,
        "expected_page_count": expected_pages,
        "selected": selected_entries,
    }
    report = {
        "schema_version": SCHEMA_VERSION,
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "software": {"name": software_name, "version": version},
        "output_docx": str(output),
        "status": "blocked" if blockers else ("dry_run" if args.dry_run else "generated_requires_office_render"),
        "submission_ready": False,
        "office_render_required": True,
        "counts": {
            "source_files": len(file_infos),
            "full_source_lines": len(full_lines),
            "deposit_lines": len(selected),
            "expected_pages": expected_pages,
            "explicit_page_breaks": max(expected_pages - 1, 0),
        },
        "selection_mode": selection_mode,
        "boundaries": {
            "first": selected[0].location(),
            "last": selected[-1].location(),
            "front_end": selected[min(lines_per_page * front_pages, len(selected)) - 1].location(),
            "back_start": (
                selected[lines_per_page * front_pages].location()
                if selection_mode == "front_and_back"
                else None
            ),
        },
        "line_length": {
            "maximum": max(line_lengths),
            "over_warning_threshold": sum(length > long_line_threshold for length in line_lengths),
            "warning_threshold": long_line_threshold,
        },
        "repeated_blocks": repeats,
        "sensitive_findings": sensitive,
        "login_risks": login_risks,
        "forbidden_selected_paths": forbidden,
        "blockers": blockers,
        "warnings": warnings,
        "next_step": "使用Word/WPS更新全部域并重新分页，导出PDF后验证实际页数和逐页页码",
    }
    atomic_write_json(provenance_output, provenance)
    atomic_write_json(report_output, report)
    if blockers:
        print(json.dumps({"status": "blocked", "report": str(report_output), "blockers": blockers}, ensure_ascii=False))
        return 3
    if not args.dry_run:
        build_code_docx(
            selected,
            output,
            software_name=software_name,
            version=version,
            deposit_config=deposit,
            expected_pages=expected_pages,
        )
    print(
        json.dumps(
            {
                "status": report["status"],
                "output": None if args.dry_run else str(output),
                "provenance": str(provenance_output),
                "report": str(report_output),
                "full_source_lines": len(full_lines),
                "deposit_lines": len(selected),
                "expected_pages": expected_pages,
            },
            ensure_ascii=False,
        )
    )
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="从JSON manifest生成可追溯的一般交存源程序DOCX（不足3000条时全量交存）。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
最小manifest示例：
{
  "project_root": "D:/project",
  "software": {"name": "示例管理系统", "version": "V1.0"},
  "source": {
    "files": ["backend/app.py", {"path": "frontend/src/App.vue", "role": "ui"}],
    "encoding_candidates": ["utf-8-sig", "utf-8", "gb18030"]
  },
  "deposit": {"lines_per_page": 50, "front_pages": 30, "back_pages": 30},
  "audit": {"repeat_widths": [5, 9, 10, 20]}
}

manifest使用“-”时从标准输入读取。显式source.files顺序就是规范源码顺序；
未提供files时可用source.include_globs，或按内置源码扩展名确定性扫描。
DOCX生成后仍须由Word/WPS更新域、重新分页并以PDF复核实际页数。
""",
    )
    parser.add_argument("--manifest", required=True, help="JSON manifest路径，或-表示标准输入")
    parser.add_argument("--output", help="覆盖manifest中的DOCX输出路径")
    parser.add_argument("--provenance", help="来源追踪JSON输出路径")
    parser.add_argument("--report", help="重复、敏感信息和登录风险报告JSON路径")
    parser.add_argument("--dry-run", action="store_true", help="只扫描并生成JSON，不生成DOCX")
    parser.add_argument(
        "--allow-sensitive",
        action="store_true",
        help="即使发现高风险敏感信息候选也继续生成；正式材料不建议使用",
    )
    parser.add_argument(
        "--fail-on-login-risk",
        action="store_true",
        help="高风险通用登录表现层进入摘录时阻止生成",
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        return run_build(args)
    except (DepositError, OSError, ValueError) as exc:
        print(json.dumps({"status": "error", "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
