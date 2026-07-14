#!/usr/bin/env python3
"""Audit deposit-source repetition and optional similarity to prior code.

The script consumes the same JSON manifest as ``build_code_deposit.py``.  It
can compare the selected deposit (default) or the full canonical snapshot with
one or more old DOCX files, source files, source directories, or JSON source
manifests.  Results are evidence for review, not a search of a registration
authority's private corpus.
"""

from __future__ import annotations

import argparse
import difflib
import hashlib
import json
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Sequence

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from build_code_deposit import (  # noqa: E402
    DEFAULT_ENCODINGS,
    DEFAULT_EXCLUDED_DIRS,
    DEFAULT_EXTENSIONS,
    DepositError,
    SourceLine,
    _as_dict,
    _as_string_list,
    analyze_repeated_blocks,
    atomic_write_json,
    decode_source_bytes,
    load_json_manifest,
    load_project_source,
    normalized_line,
    resolve_from_base,
    scan_login_risk,
    scan_sensitive_lines,
    select_deposit_lines,
    snapshot_hash,
)


SCHEMA_VERSION = "1.0"
RISK_ORDER = {"none": 0, "info": 1, "low": 2, "medium": 3, "high": 4, "critical": 5}


@dataclass(frozen=True)
class AuditLine:
    text: str
    path: str
    source_line: int
    canonical_index: int | None = None

    def location(self) -> dict[str, Any]:
        result: dict[str, Any] = {"path": self.path, "source_line": self.source_line}
        if self.canonical_index is not None:
            result["canonical_index"] = self.canonical_index
        return result


@dataclass(frozen=True)
class LoadedReference:
    label: str
    source: str
    kind: str
    lines: tuple[AuditLine, ...]
    files: int


def _audit_lines(lines: Sequence[SourceLine]) -> list[AuditLine]:
    return [
        AuditLine(
            text=item.text,
            path=item.path,
            source_line=item.source_line,
            canonical_index=item.canonical_index,
        )
        for item in lines
    ]


def _split_effective_lines(text: str, *, path: str, tab_size: int = 4) -> list[AuditLine]:
    result: list[AuditLine] = []
    for number, raw in enumerate(text.splitlines(), start=1):
        rendered = raw.expandtabs(tab_size).rstrip()
        if rendered.strip():
            result.append(AuditLine(text=rendered, path=path, source_line=number))
    return result


def _load_docx(path: Path, label: str) -> LoadedReference:
    try:
        document = Document(path)
    except Exception as exc:  # python-docx raises several package/XML exception types
        raise DepositError(f"无法读取参考DOCX {path}：{exc}") from exc
    lines: list[AuditLine] = []
    logical_line = 0
    for paragraph in document.paragraphs:
        for value in paragraph.text.splitlines() or [paragraph.text]:
            rendered = value.rstrip()
            if rendered.strip():
                logical_line += 1
                lines.append(AuditLine(rendered, path.name, logical_line))
    seen_cells: set[int] = set()
    for table_index, table in enumerate(document.tables, start=1):
        for row in table.rows:
            for cell in row.cells:
                marker = id(cell._tc)
                if marker in seen_cells:
                    continue
                seen_cells.add(marker)
                for paragraph in cell.paragraphs:
                    for value in paragraph.text.splitlines() or [paragraph.text]:
                        rendered = value.rstrip()
                        if rendered.strip():
                            logical_line += 1
                            lines.append(
                                AuditLine(
                                    rendered,
                                    f"{path.name}#table{table_index}",
                                    logical_line,
                                )
                            )
    if not lines:
        raise DepositError(f"参考DOCX没有可比较的非空文本：{path}")
    return LoadedReference(label, str(path), "docx", tuple(lines), 1)


def _directory_file_is_excluded(path: Path, root: Path) -> bool:
    try:
        relative = path.relative_to(root)
    except ValueError:
        return True
    return bool({part.lower() for part in relative.parts} & DEFAULT_EXCLUDED_DIRS)


def _strict_source_reference(
    path: Path,
    *,
    label: str,
    declared_encoding: str | None,
    encodings: Sequence[str],
    max_file_bytes: int,
) -> LoadedReference:
    raw = path.read_bytes()
    if len(raw) > max_file_bytes:
        raise DepositError(f"参考源码文件超过大小上限：{path}")
    text, _ = decode_source_bytes(
        raw,
        path=path,
        declared_encoding=declared_encoding,
        candidates=encodings,
    )
    lines = _split_effective_lines(text, path=path.name)
    if not lines:
        raise DepositError(f"参考源码没有非空行：{path}")
    return LoadedReference(label, str(path), "source_file", tuple(lines), 1)


def _load_source_directory(
    path: Path,
    *,
    label: str,
    encodings: Sequence[str],
    extensions: set[str],
    max_file_bytes: int,
) -> LoadedReference:
    lines: list[AuditLine] = []
    count = 0
    for file_path in sorted(path.rglob("*"), key=lambda item: item.as_posix().lower()):
        if not file_path.is_file() or file_path.suffix.lower() not in extensions:
            continue
        if _directory_file_is_excluded(file_path, path):
            continue
        raw = file_path.read_bytes()
        if len(raw) > max_file_bytes:
            raise DepositError(f"参考源码文件超过大小上限：{file_path}")
        text, _ = decode_source_bytes(raw, path=file_path, declared_encoding=None, candidates=encodings)
        relative = file_path.relative_to(path).as_posix()
        lines.extend(_split_effective_lines(text, path=relative))
        count += 1
    if not lines:
        raise DepositError(f"参考源码目录没有可比较文件：{path}")
    return LoadedReference(label, str(path), "source_directory", tuple(lines), count)


def _load_reference_manifest(path: Path, label: str) -> LoadedReference:
    reference_manifest, base = load_json_manifest(str(path))
    lines, files, _ = load_project_source(reference_manifest, base)
    return LoadedReference(label, str(path), "source_manifest", tuple(_audit_lines(lines)), len(files))


def _reference_entries(manifest: dict[str, Any], cli_references: Sequence[str]) -> list[tuple[Any, bool]]:
    similarity = _as_dict(manifest.get("similarity"), "similarity")
    raw = similarity.get("references")
    if raw is None:
        raw = manifest.get("prior_code_documents", manifest.get("previous_sources", []))
    if raw is None:
        raw = []
    if not isinstance(raw, list):
        raise DepositError("similarity.references必须是数组")
    entries: list[tuple[Any, bool]] = [(item, False) for item in raw]
    entries.extend((item, True) for item in cli_references)
    return entries


def load_references(
    manifest: dict[str, Any],
    base: Path,
    cli_references: Sequence[str],
) -> list[LoadedReference]:
    similarity = _as_dict(manifest.get("similarity"), "similarity")
    encodings = tuple(_as_string_list(similarity.get("encoding_candidates"), "similarity.encoding_candidates"))
    if not encodings:
        encodings = DEFAULT_ENCODINGS
    extensions_value = similarity.get("reference_extensions", sorted(DEFAULT_EXTENSIONS))
    if not isinstance(extensions_value, list):
        raise DepositError("similarity.reference_extensions必须是数组")
    extensions = {
        str(value).lower() if str(value).startswith(".") else f".{str(value).lower()}"
        for value in extensions_value
    }
    max_file_bytes = int(similarity.get("max_reference_file_bytes", 30 * 1024 * 1024))
    references: list[LoadedReference] = []
    labels: set[str] = set()

    for raw, from_cli in _reference_entries(manifest, cli_references):
        if isinstance(raw, str):
            config: dict[str, Any] = {"path": raw}
        elif isinstance(raw, dict):
            config = dict(raw)
        else:
            raise DepositError("每个similarity.references元素必须是字符串或对象")
        value = config.get("path")
        if not isinstance(value, str) or not value.strip():
            raise DepositError("参考项缺少path")
        path_base = Path.cwd() if from_cli else base
        path = resolve_from_base(value, path_base)
        if not path.exists():
            raise DepositError(f"参考路径不存在：{path}")
        label = str(config.get("label") or path.name)
        if label in labels:
            suffix = 2
            candidate = f"{label}#{suffix}"
            while candidate in labels:
                suffix += 1
                candidate = f"{label}#{suffix}"
            label = candidate
        labels.add(label)
        kind = str(config.get("type", "auto")).lower()
        declared_encoding = config.get("encoding")
        if declared_encoding is not None and not isinstance(declared_encoding, str):
            raise DepositError(f"参考项encoding必须是字符串：{path}")
        if kind == "manifest" or (kind == "auto" and path.suffix.lower() == ".json"):
            reference = _load_reference_manifest(path, label)
        elif kind == "docx" or (kind == "auto" and path.suffix.lower() == ".docx"):
            reference = _load_docx(path, label)
        elif kind == "directory" or (kind == "auto" and path.is_dir()):
            reference = _load_source_directory(
                path,
                label=label,
                encodings=encodings,
                extensions=extensions,
                max_file_bytes=max_file_bytes,
            )
        elif path.is_file():
            reference = _strict_source_reference(
                path,
                label=label,
                declared_encoding=declared_encoding,
                encodings=encodings,
                max_file_bytes=max_file_bytes,
            )
        else:
            raise DepositError(f"不支持的参考类型：{path}")
        references.append(reference)
    return references


def _trivial_line(text: str) -> bool:
    value = normalized_line(text)
    if not value:
        return True
    if re.fullmatch(r"[{}\[\]();,.:<>/\\]+", value):
        return True
    if re.fullmatch(r"</?[A-Za-z][A-Za-z0-9:_-]*\s*/?>", value):
        return True
    if re.fullmatch(r"(?i)(?:else|try|finally|return|null|none|true|false)[:;]?", value):
        return True
    return False


def _block_is_trivial(values: Sequence[str]) -> bool:
    meaningful = [value for value in values if not _trivial_line(value)]
    if not meaningful:
        return True
    character_count = sum(len(re.sub(r"\W+", "", value, flags=re.UNICODE)) for value in meaningful)
    return character_count < max(12, len(values) * 3)


def _window_index(normalized: Sequence[str], width: int) -> dict[str, list[int]]:
    result: dict[str, list[int]] = defaultdict(list)
    if width <= 0 or len(normalized) < width:
        return result
    for index in range(len(normalized) - width + 1):
        digest = hashlib.sha256("\n".join(normalized[index : index + width]).encode("utf-8")).hexdigest()
        result[digest].append(index)
    return result


def compare_windows(
    current: Sequence[AuditLine],
    reference: Sequence[AuditLine],
    *,
    width: int,
    max_examples: int,
) -> dict[str, Any]:
    current_normalized = [normalized_line(item.text) for item in current]
    reference_normalized = [normalized_line(item.text) for item in reference]
    current_index = _window_index(current_normalized, width)
    reference_index = _window_index(reference_normalized, width)
    common = set(current_index) & set(reference_index)
    nontrivial: list[str] = []
    trivial: list[str] = []
    for digest in common:
        start = current_index[digest][0]
        values = current_normalized[start : start + width]
        (trivial if _block_is_trivial(values) else nontrivial).append(digest)
    nontrivial.sort(key=lambda value: current_index[value][0])
    trivial.sort(key=lambda value: current_index[value][0])
    examples: list[dict[str, Any]] = []
    for digest in (nontrivial + trivial)[:max_examples]:
        current_positions = current_index[digest]
        reference_positions = reference_index[digest]
        first = current_positions[0]
        examples.append(
            {
                "trivial": digest in trivial,
                "current_occurrences": [current[index].location() for index in current_positions[:8]],
                "reference_occurrences": [reference[index].location() for index in reference_positions[:8]],
                "current_occurrence_count": len(current_positions),
                "reference_occurrence_count": len(reference_positions),
                "preview": current_normalized[first : first + min(width, 12)],
            }
        )
    union_size = len(set(current_index) | set(reference_index))
    return {
        "width": width,
        "current_windows": max(len(current) - width + 1, 0),
        "reference_windows": max(len(reference) - width + 1, 0),
        "unique_matching_blocks": len(common),
        "nontrivial_matching_blocks": len(nontrivial),
        "trivial_matching_blocks": len(trivial),
        "matched_current_window_occurrences": sum(len(current_index[value]) for value in common),
        "jaccard_unique_windows": round(len(common) / union_size, 8) if union_size else 0.0,
        "examples": examples,
        "examples_truncated": len(common) > max_examples,
    }


def longest_sequence_matches(
    current: Sequence[AuditLine],
    reference: Sequence[AuditLine],
    *,
    max_examples: int = 10,
) -> dict[str, Any]:
    current_normalized = [normalized_line(item.text) for item in current]
    reference_normalized = [normalized_line(item.text) for item in reference]
    matcher = difflib.SequenceMatcher(None, current_normalized, reference_normalized, autojunk=True)
    blocks = [block for block in matcher.get_matching_blocks() if block.size]
    blocks.sort(key=lambda block: (-block.size, block.a, block.b))
    examples: list[dict[str, Any]] = []
    for block in blocks[:max_examples]:
        values = current_normalized[block.a : block.a + block.size]
        examples.append(
            {
                "length": block.size,
                "trivial": _block_is_trivial(values),
                "current_start": current[block.a].location(),
                "reference_start": reference[block.b].location(),
                "preview": values[: min(block.size, 12)],
            }
        )
    nontrivial_lengths = [item["length"] for item in examples if not item["trivial"]]
    return {
        "longest_exact_contiguous_lines": blocks[0].size if blocks else 0,
        "longest_nontrivial_lines_in_examples": max(nontrivial_lengths, default=0),
        "examples": examples,
        "note": "SequenceMatcher用于定位长连续块；窗口哈希统计是主要判定依据。",
    }


def _raise_risk(current: str, candidate: str) -> str:
    return candidate if RISK_ORDER[candidate] > RISK_ORDER[current] else current


def _risk_from_cross(result_sets: Sequence[dict[str, Any]]) -> str:
    risk = "none"
    for result in result_sets:
        if not result["nontrivial_matching_blocks"]:
            continue
        width = int(result["width"])
        if width >= 20:
            risk = _raise_risk(risk, "critical")
        elif width >= 10:
            risk = _raise_risk(risk, "high")
        elif width >= 5:
            risk = _raise_risk(risk, "medium")
        else:
            risk = _raise_risk(risk, "low")
    return risk


def _risk_from_internal(results: Sequence[dict[str, Any]]) -> str:
    risk = "none"
    for result in results:
        if not result["unique_duplicate_blocks"]:
            continue
        width = int(result["width"])
        if width >= 20:
            risk = _raise_risk(risk, "high")
        elif width >= 10:
            risk = _raise_risk(risk, "medium")
        elif width >= 5:
            risk = _raise_risk(risk, "low")
        else:
            risk = _raise_risk(risk, "info")
    return risk


def _parse_widths(value: str | None, manifest: dict[str, Any]) -> list[int]:
    similarity = _as_dict(manifest.get("similarity"), "similarity")
    raw: Any
    if value:
        try:
            raw = [int(item.strip()) for item in value.split(",") if item.strip()]
        except ValueError as exc:
            raise DepositError("--widths必须是逗号分隔的整数") from exc
    else:
        raw = similarity.get("widths", [3, 5, 10, 20])
    if not isinstance(raw, list):
        raise DepositError("similarity.widths必须是整数数组")
    widths = sorted(set(int(item) for item in raw))
    if not widths or any(item < 2 or item > 100 for item in widths):
        raise DepositError("相似度窗口宽度必须在2—100之间")
    return widths


def run_audit(args: argparse.Namespace) -> int:
    manifest, base = load_json_manifest(args.manifest)
    full_lines, source_files, project_root = load_project_source(manifest, base)
    deposit = _as_dict(manifest.get("deposit"), "deposit")
    lines_per_page = int(deposit.get("lines_per_page", 50))
    front_pages = int(deposit.get("front_pages", 30))
    back_pages = int(deposit.get("back_pages", 30))
    selected, selection_mode = select_deposit_lines(
        full_lines,
        lines_per_page=lines_per_page,
        front_pages=front_pages,
        back_pages=back_pages,
    )
    current_source = full_lines if args.scope == "full" else selected
    current = _audit_lines(current_source)
    widths = _parse_widths(args.widths, manifest)
    max_examples = args.max_examples
    if max_examples < 1 or max_examples > 200:
        raise DepositError("--max-examples必须在1—200之间")

    internal = analyze_repeated_blocks(current_source, widths, max_examples=max_examples)
    references = load_references(manifest, base, args.reference or [])
    comparison_reports: list[dict[str, Any]] = []
    overall_risk = _risk_from_internal(internal)
    for reference in references:
        windows = [
            compare_windows(
                current,
                reference.lines,
                width=width,
                max_examples=max_examples,
            )
            for width in widths
        ]
        reference_risk = _risk_from_cross(windows)
        overall_risk = _raise_risk(overall_risk, reference_risk)
        comparison_reports.append(
            {
                "label": reference.label,
                "source": reference.source,
                "kind": reference.kind,
                "files": reference.files,
                "reference_lines": len(reference.lines),
                "risk_level": reference_risk,
                "window_comparisons": windows,
                "longest_sequences": longest_sequence_matches(
                    current, reference.lines, max_examples=min(max_examples, 10)
                ),
            }
        )

    software = _as_dict(manifest.get("software"), "software")
    name = software.get(
        "name", software.get("full_name", software.get("expected_name", manifest.get("software_name", "")))
    )
    version = software.get("version", software.get("expected_version", manifest.get("version", "")))
    report = {
        "schema_version": SCHEMA_VERSION,
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "software": {"name": name, "version": version},
        "project_root": str(project_root),
        "scope": args.scope,
        "selection_mode": selection_mode,
        "source_files": len(source_files),
        "full_source_lines": len(full_lines),
        "audited_lines": len(current),
        "audited_text_sha256": snapshot_hash(current_source, include_provenance=False),
        "widths": widths,
        "risk_level": overall_risk,
        "internal_repetition": {
            "risk_level": _risk_from_internal(internal),
            "blocks": internal,
            "interpretation": "内部重复可能是正常查询、模型字段或错误处理；需结合来源人工判断。",
        },
        "reference_comparisons": comparison_reports,
        "selected_sensitive_findings": scan_sensitive_lines(selected),
        "selected_login_risks": scan_login_risk(selected),
        "limitations": [
            "仅比较manifest和--reference提供的材料，未接入登记机构或互联网私有代码库。",
            "短HTML闭合标签、括号、导入语句等通用结构可能形成低价值匹配。",
            "高相似度是人工复核信号，不应通过随机改名、插入无意义代码等方式规避。",
        ],
        "reference_count": len(references),
        "warnings": [] if references else ["未提供旧DOCX或旧源码，只完成交存件内部重复审计。"],
    }
    similarity = _as_dict(manifest.get("similarity"), "similarity")
    configured_output = similarity.get("output_report")
    default_output = base / "代码相似度审计报告.json"
    output = resolve_from_base(args.output or configured_output, base) if (args.output or configured_output) else default_output.resolve()
    atomic_write_json(output, report)
    print(
        json.dumps(
            {
                "status": "audited",
                "output": str(output),
                "scope": args.scope,
                "audited_lines": len(current),
                "references": len(references),
                "risk_level": overall_risk,
            },
            ensure_ascii=False,
        )
    )
    if args.fail_on != "none" and RISK_ORDER[overall_risk] >= RISK_ORDER[args.fail_on]:
        return 3
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="审计一般交存代码的内部重复，并与可选旧DOCX、旧源码或旧manifest比较。",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python audit_similarity.py --manifest project.json
  python audit_similarity.py --manifest project.json --reference old-code.docx
  python audit_similarity.py --manifest project.json --reference D:/old/source --widths 3,5,10,20

manifest也可配置：
  "similarity": {
    "references": [
      "old-code.docx",
      {"path": "../old-project", "type": "directory", "label": "旧项目源码"}
    ],
    "widths": [3, 5, 10, 20]
  }

该工具不搜索登记机构的未知代码库，不能给出“绝对无雷同”保证。
""",
    )
    parser.add_argument("--manifest", required=True, help="与代码交存生成器共用的JSON manifest，或-")
    parser.add_argument(
        "--reference",
        action="append",
        default=[],
        help="旧DOCX、旧源码文件、源码目录或JSON manifest；可重复指定",
    )
    parser.add_argument("--output", help="审计JSON输出路径")
    parser.add_argument("--scope", choices=("deposit", "full"), default="deposit", help="审计60页摘录或完整源码")
    parser.add_argument("--widths", help="连续窗口宽度，例如3,5,10,20")
    parser.add_argument("--max-examples", type=int, default=20, help="每种宽度最多保留的示例数，默认20")
    parser.add_argument(
        "--fail-on",
        choices=("none", "medium", "high", "critical"),
        default="none",
        help="风险达到指定等级时返回退出码3；默认仅报告",
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        return run_audit(args)
    except (DepositError, OSError, ValueError) as exc:
        print(json.dumps({"status": "error", "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
